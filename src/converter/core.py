#!/usr/bin/env python3
"""
Batch2MD v7.2 — Conversor em lote para Markdown (GPU-Accelerated, NO LLM)

═══════════════════════════════════════════════════════════════════════════════
                    OTIMIZADO PARA GOOGLE COLAB A100/V100/T4
═══════════════════════════════════════════════════════════════════════════════

VERSÃO v7.2 (2025-12-17):
- NOVO: Header/footer exclusion automático via pymupdf-layout (default: excluir)
- NOVO: Detecção automática de PPT→PDF com otimizações (Issue #78)
- NOVO: Detecção MPS (Apple Silicon) com fallback para CPU
- NOVO: Parâmetros configuráveis: exclude_headers, exclude_footers, ignore_graphics
- NOVO: table_strategy configurável (lines_strict, lines, text)
- NOVO: CLI options: --exclude-headers, --include-footers, --ignore-graphics
- MELHORADO: Logging detalhado do status de layout extension
- MELHORADO: Exibição de features do pymupdf-layout na inicialização
- CORREÇÃO: PPT→PDF multi-column extraction issues (GitHub Issue #78)
- REF: https://pymupdf.readthedocs.io/en/latest/pymupdf4llm/api.html

VERSÃO v7.1 (2025-12-15):
- NOVO: Integração pymupdf-layout para extração layout-aware
- NOVO: Auto-instalação de pymupdf-layout (se ausente)
- NOVO: Pipeline aprimorado de normalização Unicode
- NOVO: Correção de fragmentação de caracteres (fix_fragmented_text)
- NOVO: Validação de coerência de texto (validate_text_coherence)
- NOVO: Flags de extração otimizados (TEXT_PRESERVE_WHITESPACE, TEXT_DEHYPHENATE)
- CORRIGIDO: Caracteres estranhos em números e valores
- CORRIGIDO: Acentos portugueses fragmentados
- MELHORADO: Ordem de leitura em documentos multi-coluna

VERSÃO v7.0 (2025-12-14):
- NOVO: Substituição completa do LLM por OCR tradicional GPU-accelerado
- NOVO: PaddleOCR PP-OCRv5 como engine principal (2MB vs 14GB do Qwen2-VL)
- NOVO: Docling TableFormer para extração de tabelas (97.9% accuracy)
- NOVO: Pipeline de preprocessamento de imagem (deskew, binarização, denoise)
- NOVO: Suporte a 5 níveis de ZIP aninhados (era 3)
- NOVO: Detecção de ZIP bombs e proteção path traversal
- NOVO: Excel multi-sheet com seção por aba
- NOVO: Cascata de fallback automático (OCR e Tables)

COMPARAÇÃO v6 vs v7:
┌─────────────────┬──────────────┬──────────────┐
│ Métrica         │ v6 (LLM)     │ v7 (OCR)     │
├─────────────────┼──────────────┼──────────────┤
│ Modelo          │ 14GB         │ 2MB          │
│ Cold Start      │ 2-5 min      │ 5-10 sec     │
│ VRAM            │ 20-40GB      │ 2-4GB        │
│ Tabelas         │ ~85%         │ 97.9%        │
│ Português       │ Bom          │ Nativo       │
└─────────────────┴──────────────┴──────────────┘

OCR ENGINE STACK:
- [1] PaddleOCR PP-OCRv5 (GPU) - 90.67% OmniBenchDoc score
- [2] EasyOCR (GPU) - Fallback PyTorch
- [3] Tesseract (CPU) - Fallback universal

TABLE EXTRACTION STACK:
- [1] Docling TableFormer - 97.9% accuracy
- [2] PyMuPDF find_tables() - Rápido, nativo
- [3] PDFPlumber "text" - Tabelas sem bordas
- [4] Camelot lattice - Tabelas com bordas

AMBIENTE ALVO:
- Google Colab A100-80GB (150GB RAM, 80GB GPU) - MODO PADRÃO
- Google Colab A100-40GB (85GB RAM, 40GB GPU) - MODO TURBO
- Google Colab V100 (26GB RAM, 16GB GPU)
- Google Colab T4 (12GB RAM, 16GB GPU) - MODO LITE
- CUDA 11.8/12.1 com PaddlePaddle

INSTALAÇÃO COLAB (A100-80GB - TESTADO 2024-12):
    # ══════════════════════════════════════════════════════════════════════
    # CELL 1: FIX NUMPY (CRÍTICO - executar primeiro!)
    # ══════════════════════════════════════════════════════════════════════
    !pip uninstall -y numpy
    !pip install "numpy<2.0"

    # ══════════════════════════════════════════════════════════════════════
    # CELL 2: Verificar GPU
    # ══════════════════════════════════════════════════════════════════════
    !nvidia-smi
    import torch
    print(f"CUDA: {torch.cuda.is_available()}")
    print(f"GPU: {torch.cuda.get_device_name(0)}")
    print(f"VRAM: {torch.cuda.get_device_properties(0).total_memory / 1e9:.1f} GB")

    # ══════════════════════════════════════════════════════════════════════
    # CELL 3: Reinstalar PaddlePaddle (limpo)
    # ══════════════════════════════════════════════════════════════════════
    !pip uninstall -y paddlepaddle paddlepaddle-gpu paddle
    !pip install paddlepaddle-gpu==2.6.1.post120 -f https://www.paddlepaddle.org.cn/whl/linux/mkl/avx/stable.html
    !pip install paddleocr==2.7.3

    # ══════════════════════════════════════════════════════════════════════
    # CELL 4: Document Processing
    # ══════════════════════════════════════════════════════════════════════
    !pip install -q pymupdf==1.24.10 pymupdf4llm pdfplumber opencv-python-headless
    !pip install -q openpyxl xlrd python-docx python-pptx extract-msg

    # ══════════════════════════════════════════════════════════════════════
    # CELL 5: OCR Fallbacks + Utils
    # ══════════════════════════════════════════════════════════════════════
    !pip install -q easyocr pytesseract typer rich markitdown deskew scikit-image

    # ══════════════════════════════════════════════════════════════════════
    # CELL 6: Tesseract engine
    # ══════════════════════════════════════════════════════════════════════
    !apt-get install -qq tesseract-ocr tesseract-ocr-por tesseract-ocr-eng

    # ══════════════════════════════════════════════════════════════════════
    # CELL 7: Reiniciar runtime (IMPORTANTE após instalar numpy!)
    # ══════════════════════════════════════════════════════════════════════
    # Menu: Runtime → Restart runtime

    # ══════════════════════════════════════════════════════════════════════
    # CELL 8: Executar (APÓS reiniciar)
    # ══════════════════════════════════════════════════════════════════════
    !python Batch2MD-v7.py input.zip --out ./md_output --verbose

USO:
    python Batch2MD-v7.py INPUT --out OUTPUT [OPTIONS]

    # Exemplos:
    python Batch2MD-v7.py arquivo.zip --out ./saida
    python Batch2MD-v7.py pasta/ --out ./saida --verbose
    python Batch2MD-v7.py doc.pdf --out ./saida
"""

from __future__ import annotations

import base64
import hashlib
import io
import os
import queue
import re
import subprocess
import sys
import html as html_module  # For HTML entity decoding
import shutil
import tempfile
import threading
import unicodedata
import zipfile
from abc import ABC, abstractmethod
from collections import deque
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Callable, Optional, Union

import typer
from rich.console import Console
from rich.progress import BarColumn, Progress, SpinnerColumn, TextColumn, TimeElapsedColumn

# ============================================
# DETECÇÃO DE AMBIENTE
# ============================================

def detect_environment() -> str:
    """Detecta ambiente de execução: colab, jupyter, local."""
    try:
        import google.colab  # noqa: F401
        return "colab"
    except ImportError:
        pass
    try:
        shell = get_ipython().__class__.__name__  # type: ignore[name-defined]
        if shell == "ZMQInteractiveShell":
            return "jupyter"
    except NameError:
        pass
    return "local"


ENVIRONMENT = detect_environment()

# ============================================
# IMPORTS OPCIONAIS COM FLAGS
# ============================================

# NumPy (required)
try:
    import numpy as np
    HAS_NUMPY = True
    # Check numpy version for ABI compatibility
    NUMPY_VERSION = tuple(map(int, np.__version__.split('.')[:2]))
    NUMPY_IS_V2 = NUMPY_VERSION >= (2, 0)
except ImportError:
    HAS_NUMPY = False
    np = None  # type: ignore
    NUMPY_VERSION = (0, 0)
    NUMPY_IS_V2 = False

# PyTorch
try:
    import torch
    HAS_TORCH = True
except ImportError:
    HAS_TORCH = False
    torch = None  # type: ignore

# OpenCV
try:
    import cv2
    HAS_OPENCV = True
    # Check for CUDA support
    HAS_OPENCV_CUDA = cv2.cuda.getCudaEnabledDeviceCount() > 0 if hasattr(cv2, 'cuda') else False
except ImportError:
    HAS_OPENCV = False
    HAS_OPENCV_CUDA = False
    cv2 = None  # type: ignore

# PIL
try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
    Image = None  # type: ignore

# ============================================
# PyMuPDF + Layout Extension (ORDEM DE IMPORT CRÍTICA!)
# pymupdf.layout DEVE ser importado ANTES de pymupdf4llm
# ============================================

def _ensure_pymupdf_layout() -> bool:
    """
    Tenta garantir que pymupdf-layout está disponível.
    Auto-instala se ausente e possível.
    """
    try:
        import pymupdf.layout  # noqa: F401
        return True
    except ImportError:
        try:
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", "-q", "pymupdf-layout"],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL
            )
            import pymupdf.layout  # noqa: F401
            return True
        except Exception:
            return False

# PyMuPDF
try:
    import pymupdf as fitz
    HAS_PYMUPDF = True
except ImportError:
    try:
        import fitz
        HAS_PYMUPDF = True
    except ImportError:
        HAS_PYMUPDF = False
        fitz = None  # type: ignore

# PyMuPDF Layout Extension (CRÍTICO: importar ANTES de pymupdf4llm)
HAS_PYMUPDF_LAYOUT = False
if HAS_PYMUPDF:
    HAS_PYMUPDF_LAYOUT = _ensure_pymupdf_layout()

# PyMuPDF4LLM (DEVE vir DEPOIS de pymupdf.layout)
try:
    import pymupdf4llm
    HAS_PYMUPDF4LLM = True
except ImportError:
    HAS_PYMUPDF4LLM = False

# PDFPlumber
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

# PaddleOCR
try:
    from paddleocr import PaddleOCR
    HAS_PADDLEOCR = True
except ImportError:
    HAS_PADDLEOCR = False
    PaddleOCR = None  # type: ignore

# EasyOCR
try:
    import easyocr
    HAS_EASYOCR = True
except ImportError:
    HAS_EASYOCR = False
    easyocr = None  # type: ignore

# Tesseract
try:
    import pytesseract
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False

# Docling
try:
    from docling.document_converter import DocumentConverter as DoclingConverter
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    HAS_DOCLING = True
except ImportError:
    HAS_DOCLING = False
    DoclingConverter = None  # type: ignore

# Camelot
try:
    import camelot
    HAS_CAMELOT = True
except ImportError:
    HAS_CAMELOT = False

# OpenPyXL
try:
    from openpyxl import load_workbook
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

# xlrd (legacy .xls)
try:
    import xlrd
    HAS_XLRD = True
except ImportError:
    HAS_XLRD = False

# python-docx
try:
    from docx import Document as DocxDocument
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

# python-pptx
try:
    from pptx import Presentation
    HAS_PYTHON_PPTX = True
except ImportError:
    HAS_PYTHON_PPTX = False

# extract-msg
try:
    import extract_msg
    HAS_EXTRACT_MSG = True
except ImportError:
    HAS_EXTRACT_MSG = False

# MarkItDown
try:
    from markitdown import MarkItDown
    HAS_MARKITDOWN = True
except ImportError:
    HAS_MARKITDOWN = False

# Deskew
try:
    from deskew import determine_skew
    HAS_DESKEW = True
except ImportError:
    HAS_DESKEW = False

# scikit-image (for Sauvola binarization)
try:
    from skimage.filters import threshold_sauvola
    HAS_SKIMAGE = True
except ImportError:
    HAS_SKIMAGE = False

# Pandas (for Excel)
try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

# ============================================
# CONSTANTES E CONFIGURAÇÃO
# ============================================

app = typer.Typer(
    add_completion=False,
    help="Batch2MD v7.1 - Conversor GPU-Accelerated para Markdown (NO LLM)"
)
console = Console()

# Extensões suportadas
SUPPORTED_EXTENSIONS: set[str] = {
    ".pdf", ".docx", ".doc", ".xlsx", ".xls", ".pptx", ".ppt",
    ".msg", ".eml", ".html", ".htm", ".xml", ".json", ".csv",
    ".jpg", ".jpeg", ".png", ".gif", ".tiff", ".bmp", ".webp",
    ".txt", ".md",
}

LEGACY_TO_MODERN: dict[str, str] = {
    ".doc": ".docx",
    ".xls": ".xlsx",
    ".ppt": ".pptx",
}

# Limites de segurança
MAX_FILE_SIZE_MB = 500
MAX_TOTAL_FILES = 10000
MAX_ZIP_DEPTH = 5  # Aumentado de 3 para 5
MAX_ZIP_RATIO = 100  # Detecção de ZIP bomb
CONVERSION_TIMEOUT = 600

# ============================================
# GPU DETECTION E CONFIGURAÇÃO
# ============================================

@dataclass
class GPUConfig:
    """Configuração otimizada por tipo de GPU."""
    gpu_type: str
    gpu_name: str
    vram_gb: float
    paddle_batch_size: int
    docling_batch_size: int
    preprocessing_workers: int
    memory_fraction: float
    is_available: bool


def detect_gpu() -> GPUConfig:
    """Detecta GPU e retorna configuração otimizada."""
    if not HAS_TORCH:
        return GPUConfig(
            gpu_type="cpu",
            gpu_name="CPU Only",
            vram_gb=0,
            paddle_batch_size=1,
            docling_batch_size=1,
            preprocessing_workers=4,
            memory_fraction=0.0,
            is_available=False,
        )

    if not torch.cuda.is_available():
        return GPUConfig(
            gpu_type="cpu",
            gpu_name="CUDA Not Available",
            vram_gb=0,
            paddle_batch_size=1,
            docling_batch_size=1,
            preprocessing_workers=4,
            memory_fraction=0.0,
            is_available=False,
        )

    try:
        props = torch.cuda.get_device_properties(0)
        vram_gb = props.total_memory / (1024**3)
        gpu_name = props.name

        # Classificação por tipo - OTIMIZADO PARA MÁXIMA PERFORMANCE
        if "A100" in gpu_name and vram_gb > 75:
            # A100-80GB: MODO ULTRA - máxima utilização
            return GPUConfig(
                gpu_type="a100_80gb",
                gpu_name=gpu_name,
                vram_gb=vram_gb,
                paddle_batch_size=128,  # Aumentado de 64
                docling_batch_size=128,  # Aumentado de 64
                preprocessing_workers=16,  # Aumentado de 8
                memory_fraction=0.95,  # Aumentado de 0.90 - usa 76GB
                is_available=True,
            )
        elif "A100" in gpu_name:
            # A100-40GB: MODO TURBO
            return GPUConfig(
                gpu_type="a100_40gb",
                gpu_name=gpu_name,
                vram_gb=vram_gb,
                paddle_batch_size=64,
                docling_batch_size=64,
                preprocessing_workers=12,
                memory_fraction=0.90,
                is_available=True,
            )
        elif "V100" in gpu_name:
            return GPUConfig(
                gpu_type="v100",
                gpu_name=gpu_name,
                vram_gb=vram_gb,
                paddle_batch_size=32,
                docling_batch_size=32,
                preprocessing_workers=8,
                memory_fraction=0.85,
                is_available=True,
            )
        elif "T4" in gpu_name:
            return GPUConfig(
                gpu_type="t4",
                gpu_name=gpu_name,
                vram_gb=vram_gb,
                paddle_batch_size=16,
                docling_batch_size=16,
                preprocessing_workers=4,
                memory_fraction=0.80,
                is_available=True,
            )
        else:
            # GPU genérica - assume capacidade média
            return GPUConfig(
                gpu_type="generic",
                gpu_name=gpu_name,
                vram_gb=vram_gb,
                paddle_batch_size=max(8, int(vram_gb / 2)),  # Escala com VRAM
                docling_batch_size=max(8, int(vram_gb / 2)),
                preprocessing_workers=8,
                memory_fraction=0.80,
                is_available=True,
            )
    except Exception:
        return GPUConfig(
            gpu_type="cpu",
            gpu_name="GPU Detection Failed",
            vram_gb=0,
            paddle_batch_size=1,
            docling_batch_size=1,
            preprocessing_workers=4,
            memory_fraction=0.0,
            is_available=False,
        )


# Detecta GPU no startup
GPU_CONFIG = detect_gpu()


# ============================================
# DETECÇÃO MPS (APPLE SILICON) - v7.2
# ============================================

def detect_mps_and_warn() -> tuple[bool, str]:
    """
    Detecta MPS (Apple Silicon) e emite warning apropriado.

    PyMuPDF Layout Extension NÃO suporta MPS (Metal Performance Shaders).
    Referência: https://pymupdf.readthedocs.io/en/latest/pymupdf-layout/index.html#limitations

    Returns:
        tuple[bool, str]: (is_mps_system, warning_message)

    Comportamento:
        - Se MPS detectado: Retorna True e mensagem de warning
        - Se MPS não disponível: Retorna False e string vazia
        - Se torch não instalado: Retorna False e string vazia

    Exemplo:
        >>> is_mps, msg = detect_mps_and_warn()
        >>> if is_mps:
        ...     print(msg)  # "⚠️ MPS (Apple Silicon) detectado..."
    """
    if not HAS_TORCH:
        return False, ""

    try:
        # Verificar se MPS está disponível (macOS 12.3+ com Apple Silicon)
        if hasattr(torch.backends, 'mps') and torch.backends.mps.is_available():
            warning_msg = (
                "⚠️ MPS (Apple Silicon) detectado. "
                "PyMuPDF Layout Extension não suporta MPS - usando CPU para layout detection. "
                "Isso pode resultar em extração menos precisa de documentos multi-coluna. "
                "Ref: https://pymupdf.readthedocs.io/en/latest/pymupdf-layout/index.html#limitations"
            )
            console.print(f"[yellow]{warning_msg}[/yellow]")
            return True, warning_msg
    except Exception:
        # Silencioso em caso de erro na detecção
        pass

    return False, ""


# Executar detecção MPS na inicialização do módulo
IS_MPS_SYSTEM, MPS_WARNING_MESSAGE = detect_mps_and_warn()


# ============================================
# CONFIGURAÇÃO DE CONVERSÃO
# ============================================

@dataclass
class ConversionConfig:
    """Configuração completa para conversão v7."""

    # Processing
    recursive: bool = True
    workers: int = 8

    # ZIP
    process_nested_zips: bool = True
    max_zip_depth: int = MAX_ZIP_DEPTH

    # OCR
    ocr_engine: str = "auto"  # auto, paddle, easyocr, tesseract
    ocr_languages: list[str] = field(default_factory=lambda: ["pt", "en"])

    # Preprocessing (máxima qualidade por padrão)
    enable_deskew: bool = True
    enable_binarize: bool = True
    enable_denoise: bool = True
    preprocessing_dpi: int = 300

    # Tables
    table_engine: str = "auto"  # auto, docling, pymupdf, pdfplumber, camelot

    # Output
    add_frontmatter: bool = True
    add_page_markers: bool = True
    flat_names: bool = True

    # GPU
    force_cpu: bool = False
    force_gpu: bool = False  # Fail if GPU not available
    gpu_memory_fraction: float = 0.85

    # ============================================
    # Layout-aware extraction (pymupdf-layout) - v7.2 EXPANDIDO
    # ============================================
    # Referência: https://pymupdf.readthedocs.io/en/latest/pymupdf4llm/api.html

    use_layout_extraction: bool = True
    """Usar pymupdf4llm quando disponível. Se False, usa extração página-a-página."""

    force_text_over_graphics: bool = True
    """Extrair texto mesmo quando sobreposto a gráficos (force_text param)."""

    # v7.2 - Header/Footer Exclusion
    exclude_headers: bool = True
    """
    Excluir headers da extração de texto.
    REQUER: pymupdf-layout instalado (HAS_PYMUPDF_LAYOUT=True).
    Mapeia para: pymupdf4llm.to_markdown(header=False)
    """

    exclude_footers: bool = True
    """
    Excluir footers da extração de texto.
    REQUER: pymupdf-layout instalado (HAS_PYMUPDF_LAYOUT=True).
    Mapeia para: pymupdf4llm.to_markdown(footer=False)
    """

    # v7.2 - Table Strategy
    table_strategy: str = "lines_strict"
    """
    Estratégia de detecção de tabelas.
    Opções: "lines_strict" (default), "lines", "text"
    """

    # v7.2 - Graphics/Images Handling
    ignore_graphics_for_text: bool = False
    """Ignorar gráficos durante extração de texto."""

    ignore_images_for_text: bool = False
    """Ignorar imagens durante extração de texto."""

    # v7.2 - Margins Configuration
    layout_margins: tuple[float, ...] = field(default_factory=lambda: (0,))
    """Margens para exclusão de conteúdo (em pontos)."""

    # v7.2 - PPT Detection
    auto_detect_ppt_origin: bool = True
    """
    Detectar automaticamente PDFs originados de PowerPoint/Apresentações.
    Ref: https://github.com/pymupdf/RAG/issues/78
    """

    # Text quality / fragmentação
    fix_fragmentation: bool = True
    """Habilitar correção de fragmentação de caracteres."""

    validate_coherence: bool = True
    """Validar coerência do texto extraído."""

    # Debug
    verbose: bool = False
    dry_run: bool = False
    silent_mode: bool = False


# ============================================
# UTILITÁRIOS DE NORMALIZAÇÃO UTF-8
# ============================================

def normalize_unicode_text(text: str) -> str:
    """
    Normaliza texto Unicode para UTF-8 com tratamento de fragmentação.

    Pipeline otimizado:
    1. Remove zero-width chars PRIMEIRO (evita fragmentação)
    2. NFD → NFC (junta acentos combinados)
    3. NFKC (compatibilidade)
    4. Remove controle chars
    5. Normaliza espaços
    """
    if not text:
        return ""

    # PASSO 1: Remover zero-width chars ANTES de qualquer processamento
    # Esses caracteres podem causar fragmentação de texto
    zero_width_chars = [
        "\u200b",  # Zero-width space
        "\u200c",  # Zero-width non-joiner
        "\u200d",  # Zero-width joiner
        "\u200e",  # Left-to-right mark
        "\u200f",  # Right-to-left mark
        "\u2060",  # Word joiner
        "\ufeff",  # BOM / Zero-width no-break space
        "\u00ad",  # Soft hyphen
        "\u034f",  # Combining grapheme joiner
        "\u061c",  # Arabic letter mark
        "\u115f",  # Hangul choseong filler
        "\u1160",  # Hangul jungseong filler
        "\u17b4",  # Khmer vowel inherent Aq
        "\u17b5",  # Khmer vowel inherent Aa
        "\u180e",  # Mongolian vowel separator
        "\u2061",  # Function application
        "\u2062",  # Invisible times
        "\u2063",  # Invisible separator
        "\u2064",  # Invisible plus
    ]
    for char in zero_width_chars:
        text = text.replace(char, "")

    # PASSO 2: NFD → NFC para juntar acentos combinados
    # Ex: 'a' + combining acute (U+0301) → 'á' (U+00E1)
    try:
        text = unicodedata.normalize("NFD", text)
        text = unicodedata.normalize("NFC", text)
    except Exception:
        pass

    # PASSO 3: NFKC para compatibilidade
    try:
        text = unicodedata.normalize("NFKC", text)
    except Exception:
        pass

    # PASSO 4: Remove caracteres de controle (exceto newlines/tabs)
    control_chars = "".join(chr(i) for i in range(32) if chr(i) not in "\n\r\t")
    control_chars += "".join(chr(i) for i in range(127, 160))
    text = text.translate(str.maketrans("", "", control_chars))

    # PASSO 5: Normaliza variações de espaço em branco
    whitespace_map = {
        "\u00a0": " ",  # Non-breaking space
        "\u2002": " ",  # En space
        "\u2003": " ",  # Em space
        "\u2004": " ",  # Three-per-em space
        "\u2005": " ",  # Four-per-em space
        "\u2006": " ",  # Six-per-em space
        "\u2007": " ",  # Figure space
        "\u2008": " ",  # Punctuation space
        "\u2009": " ",  # Thin space
        "\u200a": " ",  # Hair space
        "\u202f": " ",  # Narrow no-break space
        "\u205f": " ",  # Medium mathematical space
        "\u3000": " ",  # Ideographic space
    }
    for old, new in whitespace_map.items():
        text = text.replace(old, new)

    # PASSO 6: Normaliza aspas e travessões
    char_replacements = {
        "\u2018": "'",  # Left single quote
        "\u2019": "'",  # Right single quote
        "\u201a": "'",  # Single low-9 quote
        "\u201b": "'",  # Single high-reversed-9 quote
        "\u201c": '"',  # Left double quote
        "\u201d": '"',  # Right double quote
        "\u201e": '"',  # Double low-9 quote
        "\u201f": '"',  # Double high-reversed-9 quote
        "\u2013": "-",  # En dash
        "\u2014": "-",  # Em dash
        "\u2015": "-",  # Horizontal bar
        "\u2026": "...",  # Ellipsis
    }
    for old, new in char_replacements.items():
        text = text.replace(old, new)

    # PASSO 7: Colapsa espaços múltiplos
    text = re.sub(r"[ \t]+", " ", text)

    # PASSO 8: Limpa linhas
    lines = text.split("\n")
    lines = [line.rstrip() for line in lines]
    text = "\n".join(lines)

    # PASSO 9: Encode/decode UTF-8 final para segurança
    try:
        text = text.encode("utf-8", errors="ignore").decode("utf-8", errors="ignore")
    except Exception:
        pass

    return text


def fix_encoding_errors(text: str) -> str:
    """Corrige mojibake (UTF-8 mal interpretado)."""
    if not text:
        return ""

    mojibake_patterns = ["Ã£", "Ã§", "Ã©", "Ã¡", "Ã³", "Ãº", "Ã¢", "Ãª", "Ã´", "Ã­"]
    has_mojibake = any(pattern in text for pattern in mojibake_patterns)

    if has_mojibake:
        try:
            fixed = text.encode("latin-1", errors="ignore").decode("utf-8", errors="ignore")
            if fixed and "�" not in fixed:
                return fixed
        except (UnicodeDecodeError, UnicodeEncodeError):
            pass

    encoding_fixes = {
        "Ã£": "ã", "Ãµ": "õ", "Ã¡": "á", "Ã©": "é", "Ã­": "í",
        "Ã³": "ó", "Ãº": "ú", "Ã¢": "â", "Ãª": "ê", "Ã´": "ô",
        "Ã§": "ç", "Ã ": "à",
    }
    for wrong, correct in sorted(encoding_fixes.items(), key=lambda x: -len(x[0])):
        text = text.replace(wrong, correct)
    return text


def remove_cid_codes(text: str) -> str:
    """Remove códigos CID de PDFs."""
    if not text or "(cid:" not in text:
        return text
    return re.sub(r"\(cid:\d+\)", "", text)


def fix_fragmented_text(text: str) -> str:
    """
    Corrige fragmentação de caracteres comum em extração de PDF.

    Problemas tratados:
    - Letras separadas por espaços: "r e c o n h e c i m e n t o"
    - Parênteses órfãos no início de palavras: "()epalavra"
    - Números fragmentados: "1 . 234 , 56" → "1.234,56"
    - Caracteres de controle residuais entre letras
    """
    if not text:
        return ""

    # Padrão 1: Letras separadas por espaços (3+ letras)
    # Ex: "r e c o n h e c i m e n t o" → "reconhecimento"
    def join_spaced_letters(match: re.Match) -> str:
        return match.group(0).replace(" ", "")

    # Regex para sequência de letras (incluindo acentos) separadas por espaço único
    spaced_pattern = r'\b([a-zA-ZáéíóúâêôãõçÁÉÍÓÚÂÊÔÃÕÇàèìòùÀÈÌÒÙ]\s){3,}[a-zA-ZáéíóúâêôãõçÁÉÍÓÚÂÊÔÃÕÇàèìòùÀÈÌÒÙ]\b'
    text = re.sub(spaced_pattern, join_spaced_letters, text)

    # Padrão 2: Parênteses órfãos seguidos de palavra
    # Ex: "()eoreconhecimento" → "eoreconhecimento"
    text = re.sub(r'\(\s*\)(?=[a-zA-ZáéíóúâêôãõçÁÉÍÓÚÂÊÔÃÕÇ])', '', text)

    # Padrão 3: Números fragmentados com pontos (formato brasileiro)
    # Ex: "1 . 234 . 567" → "1.234.567"
    text = re.sub(r'(\d)\s*\.\s*(\d)', r'\1.\2', text)

    # Padrão 4: Números fragmentados com vírgulas (decimais brasileiros)
    # Ex: "1.234 , 56" → "1.234,56"
    text = re.sub(r'(\d)\s*,\s*(\d)', r'\1,\2', text)

    # Padrão 5: Valores monetários fragmentados
    # Ex: "R $ 1.234,56" → "R$ 1.234,56"
    text = re.sub(r'R\s*\$\s*', 'R$ ', text)

    # Padrão 6: Porcentagens fragmentadas
    # Ex: "10 %" → "10%"
    text = re.sub(r'(\d)\s+%', r'\1%', text)

    # Padrão 7: Unidades fragmentadas comuns
    # Ex: "10 km" com espaço invisível → "10 km" com espaço normal
    text = re.sub(r'(\d)\s+(km|m|cm|mm|kg|g|mg|l|ml|ha|MW|kW|GW)\b', r'\1 \2', text)

    # Padrão 8: Datas fragmentadas (DD/MM/AAAA ou DD.MM.AAAA)
    # Ex: "12 / 03 / 2025" → "12/03/2025"
    text = re.sub(r'(\d{1,2})\s*/\s*(\d{1,2})\s*/\s*(\d{2,4})', r'\1/\2/\3', text)
    text = re.sub(r'(\d{1,2})\s*\.\s*(\d{1,2})\s*\.\s*(\d{2,4})', r'\1.\2.\3', text)

    # Padrão 9: Referências de processo fragmentadas
    # Ex: "50500 . 123456 / 2025 - 11" → "50500.123456/2025-11"
    text = re.sub(r'(\d{5})\s*\.\s*(\d{6})\s*/\s*(\d{4})\s*-\s*(\d{2})', r'\1.\2/\3-\4', text)

    return text


def validate_text_coherence(text: str) -> str:
    """
    Valida coerência do texto extraído e tenta corrigir problemas.

    Detecta:
    - Excesso de palavras de uma letra (indicativo de fragmentação)
    - Padrões de texto corrompido
    """
    if not text or len(text) < 50:
        return text

    words = text.split()
    if not words:
        return text

    # Conta palavras de uma única letra (exceto artigos/preposições)
    valid_single_chars = {'a', 'e', 'i', 'o', 'u', 'é', 'à', 'A', 'E', 'I', 'O', 'U', 'É', 'À'}
    single_char_words = sum(
        1 for w in words
        if len(w) == 1 and w.isalpha() and w not in valid_single_chars
    )
    single_char_ratio = single_char_words / len(words) if words else 0

    # Se mais de 20% são letras únicas suspeitas, tentar juntar
    if single_char_ratio > 0.20:
        result_words = []
        current_fragment = []

        for word in words:
            if len(word) == 1 and word.isalpha() and word not in valid_single_chars:
                current_fragment.append(word)
            else:
                if current_fragment:
                    # Junta fragmento se tem 3+ letras
                    if len(current_fragment) >= 3:
                        result_words.append("".join(current_fragment))
                    else:
                        result_words.extend(current_fragment)
                    current_fragment = []
                result_words.append(word)

        if current_fragment:
            if len(current_fragment) >= 3:
                result_words.append("".join(current_fragment))
            else:
                result_words.extend(current_fragment)

        text = " ".join(result_words)

    return text


def ensure_utf8_text(text: str) -> str:
    """
    Pipeline completo de garantia UTF-8 com correção de fragmentação.

    Ordem de processamento:
    1. fix_encoding_errors() - Corrige mojibake
    2. remove_cid_codes() - Remove placeholders de fonte
    3. normalize_unicode_text() - Normalização Unicode completa
    4. fix_fragmented_text() - Corrige fragmentação de caracteres
    5. validate_text_coherence() - Valida e corrige coerência
    """
    if not text:
        return ""

    text = fix_encoding_errors(text)
    text = remove_cid_codes(text)
    text = normalize_unicode_text(text)
    text = fix_fragmented_text(text)
    text = validate_text_coherence(text)

    return text.rstrip() + "\n" if text else ""


# ============================================
# LOGGER SIMPLIFICADO
# ============================================

class SimpleLogger:
    """Logger simplificado com contadores."""

    def __init__(self, verbose: bool = False) -> None:
        self.verbose = verbose
        self.counters: dict[str, int] = {
            "files_processed": 0,
            "zips_extracted": 0,
            "pdfs_converted": 0,
            "tables_extracted": 0,
            "ocr_pages": 0,
            "errors": 0,
        }

    def debug(self, msg: str) -> None:
        if self.verbose:
            console.print(f"[dim]{msg}[/dim]")

    def info(self, msg: str) -> None:
        console.print(f"[cyan]{msg}[/cyan]")

    def success(self, msg: str) -> None:
        console.print(f"[green]✓ {msg}[/green]")

    def warning(self, msg: str) -> None:
        console.print(f"[yellow]⚠ {msg}[/yellow]")

    def error(self, msg: str) -> None:
        self.counters["errors"] += 1
        console.print(f"[red]✗ {msg}[/red]")

    def increment(self, key: str, value: int = 1) -> None:
        self.counters[key] = self.counters.get(key, 0) + value


# ============================================
# PREPROCESSADOR DE IMAGEM (GPU-ACCELERATED)
# ============================================

class ImagePreprocessor:
    """
    Pipeline de preprocessamento de imagem para OCR.

    Etapas (máxima qualidade):
    1. Conversão para grayscale
    2. Deskewing (correção de inclinação)
    3. Binarização Sauvola (adaptativa)
    4. Redução de ruído
    5. Normalização DPI
    """

    def __init__(
        self,
        target_dpi: int = 300,
        enable_deskew: bool = True,
        enable_binarize: bool = True,
        enable_denoise: bool = True,
        use_gpu: bool = True,
    ):
        self.target_dpi = target_dpi
        self.enable_deskew = enable_deskew
        self.enable_binarize = enable_binarize
        self.enable_denoise = enable_denoise
        self.use_gpu = use_gpu and HAS_OPENCV_CUDA

    def preprocess(self, image: np.ndarray) -> np.ndarray:
        """Pipeline completo de preprocessamento."""
        if not HAS_OPENCV or not HAS_NUMPY:
            return image

        # 1. Grayscale
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image.copy()

        # 2. Deskew
        if self.enable_deskew and HAS_DESKEW:
            gray = self._deskew(gray)

        # 3. Binarização Sauvola
        if self.enable_binarize and HAS_SKIMAGE:
            gray = self._binarize_sauvola(gray)

        # 4. Denoise
        if self.enable_denoise:
            gray = self._denoise(gray)

        return gray

    def _deskew(self, image: np.ndarray) -> np.ndarray:
        """Corrige inclinação usando Hough transform."""
        try:
            angle = determine_skew(image)
            if angle is not None and abs(angle) > 0.5:
                (h, w) = image.shape[:2]
                center = (w // 2, h // 2)
                M = cv2.getRotationMatrix2D(center, angle, 1.0)
                rotated = cv2.warpAffine(
                    image, M, (w, h),
                    flags=cv2.INTER_CUBIC,
                    borderMode=cv2.BORDER_REPLICATE
                )
                return rotated
        except Exception:
            pass
        return image

    def _binarize_sauvola(self, image: np.ndarray, window_size: int = 25) -> np.ndarray:
        """Binarização adaptativa Sauvola."""
        try:
            thresh = threshold_sauvola(image, window_size=window_size)
            binary = (image > thresh).astype(np.uint8) * 255
            return binary
        except Exception:
            # Fallback para Otsu
            _, binary = cv2.threshold(image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            return binary

    def _denoise(self, image: np.ndarray) -> np.ndarray:
        """Redução de ruído."""
        try:
            # Non-local means denoising (melhor qualidade)
            return cv2.fastNlMeansDenoising(image, None, 10, 7, 21)
        except Exception:
            # Fallback para median blur
            return cv2.medianBlur(image, 3)


# ============================================
# OCR ENGINE ABSTRACT BASE
# ============================================

class OCREngine(ABC):
    """Interface abstrata para engines de OCR."""

    @property
    @abstractmethod
    def name(self) -> str:
        """Nome da engine."""
        ...

    @property
    @abstractmethod
    def supports_gpu(self) -> bool:
        """Se suporta GPU."""
        ...

    @abstractmethod
    def initialize(self) -> bool:
        """Inicializa a engine. Retorna True se sucesso."""
        ...

    @abstractmethod
    def process_image(self, image: np.ndarray) -> str:
        """Processa uma imagem e retorna texto."""
        ...

    def process_batch(self, images: list[np.ndarray]) -> list[str]:
        """Processa batch de imagens. Default: sequencial."""
        return [self.process_image(img) for img in images]


# ============================================
# PADDLEOCR ENGINE (PRIMARY)
# ============================================

class PaddleOCREngine(OCREngine):
    """
    Engine principal usando PaddleOCR PP-OCRv5.

    Performance:
    - Accuracy: 90.67 OmniBenchDoc score
    - Model size: ~2MB
    - Languages: 100+ incluindo Português nativo
    """

    name = "PaddleOCR"
    supports_gpu = True

    def __init__(
        self,
        use_gpu: bool = True,
        languages: list[str] = None,
        enable_layout: bool = True,
    ):
        self.use_gpu = use_gpu and GPU_CONFIG.is_available
        self.languages = languages or ["pt", "en"]
        self.enable_layout = enable_layout
        self._ocr = None
        self._initialized = False

    def initialize(self) -> bool:
        """Inicializa PaddleOCR com configuração explícita de GPU."""
        if not HAS_PADDLEOCR:
            return False

        try:
            # Determina lang baseado na lista
            lang = "pt" if "pt" in self.languages else "en"

            # Configura GPU explicitamente se disponível
            if self.use_gpu:
                try:
                    import paddle
                    # Força uso de GPU
                    paddle.set_device('gpu:0')
                    console.print("[green]   └─ PaddleOCR: GPU mode enabled[/green]")
                except Exception as gpu_err:
                    console.print(f"[yellow]   └─ PaddleOCR GPU setup failed: {gpu_err}, using CPU[/yellow]")
                    self.use_gpu = False

            self._ocr = PaddleOCR(
                use_angle_cls=True,
                lang=lang,
                use_gpu=self.use_gpu,
                show_log=False,
                use_space_char=True,
                gpu_mem=GPU_CONFIG.memory_fraction * 1000 if self.use_gpu else 500,  # MB
            )
            self._initialized = True
            return True
        except Exception as e:
            console.print(f"[yellow]PaddleOCR init failed: {e}[/yellow]")
            return False

    def process_image(self, image: np.ndarray) -> str:
        """Processa imagem com PaddleOCR."""
        if not self._initialized:
            if not self.initialize():
                return ""

        try:
            result = self._ocr.ocr(image, cls=True)
            if not result or not result[0]:
                return ""

            # Extrai texto mantendo ordem de leitura
            lines = []
            for line in result[0]:
                if line and len(line) >= 2:
                    text = line[1][0] if isinstance(line[1], tuple) else str(line[1])
                    lines.append(text)

            return "\n".join(lines)
        except Exception as e:
            console.print(f"[dim]PaddleOCR error: {e}[/dim]")
            return ""

    def process_batch(self, images: list[np.ndarray]) -> list[str]:
        """Processa batch de imagens para máxima performance em A100."""
        if not self._initialized:
            if not self.initialize():
                return [""] * len(images)

        results = []
        batch_size = GPU_CONFIG.paddle_batch_size

        # Processa em batches para otimizar uso da GPU
        for i in range(0, len(images), batch_size):
            batch = images[i:i + batch_size]
            batch_results = []

            for img in batch:
                try:
                    result = self._ocr.ocr(img, cls=True)
                    if result and result[0]:
                        lines = []
                        for line in result[0]:
                            if line and len(line) >= 2:
                                text = line[1][0] if isinstance(line[1], tuple) else str(line[1])
                                lines.append(text)
                        batch_results.append("\n".join(lines))
                    else:
                        batch_results.append("")
                except Exception:
                    batch_results.append("")

            results.extend(batch_results)

        return results


# ============================================
# EASYOCR ENGINE (FALLBACK 1)
# ============================================

class EasyOCREngine(OCREngine):
    """
    Engine fallback usando EasyOCR (PyTorch).

    Performance:
    - Accuracy: Alta para múltiplos idiomas
    - GPU: PyTorch nativo
    - Languages: 80+
    """

    name = "EasyOCR"
    supports_gpu = True

    def __init__(
        self,
        use_gpu: bool = True,
        languages: list[str] = None,
        batch_size: int = 8,
    ):
        self.use_gpu = use_gpu and GPU_CONFIG.is_available
        self.languages = languages or ["pt", "en"]
        self.batch_size = batch_size
        self._reader = None
        self._initialized = False

    def initialize(self) -> bool:
        """Inicializa EasyOCR com configuração explícita de GPU."""
        if not HAS_EASYOCR:
            return False

        try:
            # Força CUDA se disponível
            if self.use_gpu and HAS_TORCH:
                if torch.cuda.is_available():
                    torch.cuda.set_device(0)
                    console.print("[green]   └─ EasyOCR: CUDA device 0 selected[/green]")
                else:
                    console.print("[yellow]   └─ EasyOCR: CUDA not available, using CPU[/yellow]")
                    self.use_gpu = False

            self._reader = easyocr.Reader(
                self.languages,
                gpu=self.use_gpu,
                verbose=False,
            )
            self._initialized = True
            return True
        except Exception as e:
            console.print(f"[yellow]EasyOCR init failed: {e}[/yellow]")
            return False

    def process_image(self, image: np.ndarray) -> str:
        """Processa imagem com EasyOCR."""
        if not self._initialized:
            if not self.initialize():
                return ""

        try:
            result = self._reader.readtext(image)
            if not result:
                return ""

            lines = [item[1] for item in result if item and len(item) >= 2]
            return "\n".join(lines)
        except Exception as e:
            console.print(f"[dim]EasyOCR error: {e}[/dim]")
            return ""


# ============================================
# TESSERACT ENGINE (FALLBACK 2 - CPU)
# ============================================

class TesseractEngine(OCREngine):
    """
    Engine fallback CPU usando Tesseract.

    Sempre disponível como último recurso.
    """

    name = "Tesseract"
    supports_gpu = False

    def __init__(
        self,
        languages: list[str] = None,
        config: str = "--oem 3 --psm 6",
    ):
        # Mapeia códigos de idioma
        lang_map = {"pt": "por", "en": "eng"}
        langs = languages or ["pt", "en"]
        self.lang_str = "+".join(lang_map.get(l, l) for l in langs)
        self.config = config
        self._initialized = False

    def initialize(self) -> bool:
        """Verifica se Tesseract está disponível."""
        if not HAS_TESSERACT:
            return False

        try:
            # Testa se Tesseract está instalado
            pytesseract.get_tesseract_version()
            self._initialized = True
            return True
        except Exception as e:
            console.print(f"[yellow]Tesseract not available: {e}[/yellow]")
            return False

    def process_image(self, image: np.ndarray) -> str:
        """Processa imagem com Tesseract."""
        if not self._initialized:
            if not self.initialize():
                return ""

        try:
            # Converte para PIL se necessário
            if HAS_PIL:
                pil_image = Image.fromarray(image)
                text = pytesseract.image_to_string(
                    pil_image,
                    lang=self.lang_str,
                    config=self.config
                )
            else:
                text = pytesseract.image_to_string(
                    image,
                    lang=self.lang_str,
                    config=self.config
                )
            return text
        except Exception as e:
            console.print(f"[dim]Tesseract error: {e}[/dim]")
            return ""


# ============================================
# OCR CASCADE (FALLBACK AUTOMÁTICO)
# ============================================

class OCRCascade:
    """
    Cascata de OCR com fallback automático.

    Ordem:
    1. PaddleOCR PP-OCRv5 (GPU)
    2. EasyOCR (GPU)
    3. Tesseract (CPU)
    """

    def __init__(
        self,
        config: ConversionConfig,
        logger: SimpleLogger,
    ):
        self.config = config
        self.logger = logger
        self.engines: list[OCREngine] = []
        self._active_engine: Optional[OCREngine] = None

    def initialize(self) -> bool:
        """Inicializa engines na ordem de prioridade."""
        use_gpu = not self.config.force_cpu and GPU_CONFIG.is_available
        langs = self.config.ocr_languages

        # Determina qual engine usar
        engine_preference = self.config.ocr_engine.lower()

        if engine_preference == "auto":
            # Tenta todas em ordem
            engines_to_try = [
                PaddleOCREngine(use_gpu=use_gpu, languages=langs),
                EasyOCREngine(use_gpu=use_gpu, languages=langs),
                TesseractEngine(languages=langs),
            ]
        elif engine_preference == "paddle":
            engines_to_try = [PaddleOCREngine(use_gpu=use_gpu, languages=langs)]
        elif engine_preference == "easyocr":
            engines_to_try = [EasyOCREngine(use_gpu=use_gpu, languages=langs)]
        elif engine_preference == "tesseract":
            engines_to_try = [TesseractEngine(languages=langs)]
        else:
            engines_to_try = [
                PaddleOCREngine(use_gpu=use_gpu, languages=langs),
                TesseractEngine(languages=langs),
            ]

        # Inicializa em ordem
        for engine in engines_to_try:
            try:
                if engine.initialize():
                    self.engines.append(engine)
                    self.logger.debug(f"OCR engine initialized: {engine.name}")
            except Exception as e:
                self.logger.debug(f"OCR engine {engine.name} failed: {e}")

        if not self.engines:
            self.logger.error("No OCR engine available!")
            return False

        self._active_engine = self.engines[0]
        self.logger.info(f"OCR: {self._active_engine.name} (primary)")

        # WARN if falling back to Tesseract (CPU-only) when GPU is available
        if self._active_engine.name == "Tesseract" and GPU_CONFIG.is_available:
            console.print("\n[bold yellow]⚠ WARNING: GPU detected but using CPU-only OCR (Tesseract)[/bold yellow]")
            console.print("[yellow]  PaddleOCR/EasyOCR failed to initialize. This will be ~10x slower.[/yellow]")
            if NUMPY_IS_V2:
                console.print(f"[yellow]  ⚡ NumPy {np.__version__} (v2.x) detected - likely ABI mismatch![/yellow]")
                console.print("[yellow]  To fix, run in a fresh Colab runtime:[/yellow]")
                console.print("[dim]    !pip uninstall -y numpy[/dim]")
                console.print("[dim]    !pip install numpy==1.26.4[/dim]")
                console.print("[dim]    !pip install paddlepaddle-gpu paddleocr[/dim]")
                console.print("[dim]    # Then: Runtime → Restart session[/dim]")
            console.print("")

        return True

    def process_image(self, image: np.ndarray) -> str:
        """Processa imagem com fallback automático."""
        last_error = None

        for engine in self.engines:
            try:
                result = engine.process_image(image)
                if result and result.strip():
                    return result
            except Exception as e:
                last_error = e
                self.logger.debug(f"{engine.name} failed: {e}")
                continue

        if last_error:
            self.logger.warning(f"All OCR engines failed. Last error: {last_error}")
        return ""

    def process_batch(self, images: list[np.ndarray]) -> list[str]:
        """Processa batch de imagens."""
        return [self.process_image(img) for img in images]


# ============================================
# TABLE EXTRACTOR ABSTRACT BASE
# ============================================

class TableExtractor(ABC):
    """Interface abstrata para extratores de tabela."""

    @property
    @abstractmethod
    def name(self) -> str:
        """Nome do extrator."""
        ...

    @abstractmethod
    def initialize(self) -> bool:
        """Inicializa o extrator."""
        ...

    @abstractmethod
    def extract_tables(self, source: Any) -> list[str]:
        """Extrai tabelas e retorna como markdown."""
        ...


# ============================================
# DOCLING TABLE EXTRACTOR (PRIMARY)
# ============================================

class DoclingTableExtractor(TableExtractor):
    """
    Extrator primário usando IBM Docling TableFormer.

    Accuracy: 97.9% em tabelas complexas
    GPU: Suportado
    """

    name = "Docling"

    def __init__(self, use_gpu: bool = True):
        self.use_gpu = use_gpu and GPU_CONFIG.is_available
        self._converter = None
        self._initialized = False

    def initialize(self) -> bool:
        """Inicializa Docling."""
        if not HAS_DOCLING:
            return False

        try:
            self._converter = DoclingConverter()
            self._initialized = True
            return True
        except Exception as e:
            console.print(f"[yellow]Docling init failed: {e}[/yellow]")
            return False

    def extract_tables(self, source: Any) -> list[str]:
        """Extrai tabelas de PDF usando Docling."""
        if not self._initialized:
            if not self.initialize():
                return []

        try:
            # source pode ser Path ou string
            result = self._converter.convert(str(source))

            tables = []
            if hasattr(result, 'document') and hasattr(result.document, 'tables'):
                for table in result.document.tables:
                    if hasattr(table, 'export_to_markdown'):
                        md = table.export_to_markdown()
                        if md:
                            tables.append(md)
            return tables
        except Exception as e:
            console.print(f"[dim]Docling table extraction error: {e}[/dim]")
            return []


# ============================================
# PYMUPDF TABLE EXTRACTOR (FALLBACK 1)
# ============================================

class PyMuPDFTableExtractor(TableExtractor):
    """
    Extrator usando PyMuPDF nativo.

    Rápido, bom para tabelas bem estruturadas.
    """

    name = "PyMuPDF"

    def __init__(self):
        self._initialized = False

    def initialize(self) -> bool:
        """Verifica se PyMuPDF está disponível."""
        if not HAS_PYMUPDF:
            return False
        self._initialized = True
        return True

    def extract_tables(self, source: Any) -> list[str]:
        """Extrai tabelas de página PyMuPDF."""
        if not self._initialized:
            return []

        tables = []
        try:
            # source é fitz.Page
            if hasattr(source, 'find_tables'):
                found_tables = source.find_tables()
                if found_tables and found_tables.tables:
                    for table in found_tables.tables:
                        try:
                            md = table.to_markdown()
                            if md and md.strip():
                                tables.append(md)
                        except Exception:
                            pass
        except Exception as e:
            console.print(f"[dim]PyMuPDF table error: {e}[/dim]")
        return tables


# ============================================
# PDFPLUMBER TABLE EXTRACTOR (FALLBACK 2)
# ============================================

class PDFPlumberTableExtractor(TableExtractor):
    """
    Extrator usando PDFPlumber.

    Bom para tabelas sem bordas claras.
    """

    name = "PDFPlumber"

    def __init__(self):
        self._initialized = False

    def initialize(self) -> bool:
        """Verifica se PDFPlumber está disponível."""
        if not HAS_PDFPLUMBER:
            return False
        self._initialized = True
        return True

    def extract_tables(self, source: Any) -> list[str]:
        """Extrai tabelas de página PDFPlumber."""
        if not self._initialized:
            return []

        tables = []
        try:
            # source é pdfplumber.Page
            if hasattr(source, 'extract_tables'):
                table_settings = {
                    "vertical_strategy": "text",
                    "horizontal_strategy": "text",
                    "text_keep_blank_chars": True,
                }
                extracted = source.extract_tables(table_settings)
                for table_data in extracted:
                    if table_data:
                        md = self._table_to_markdown(table_data)
                        if md:
                            tables.append(md)
        except Exception as e:
            console.print(f"[dim]PDFPlumber table error: {e}[/dim]")
        return tables

    def _table_to_markdown(self, table: list[list]) -> str:
        """Converte tabela para markdown."""
        if not table or not table[0]:
            return ""

        lines = []
        # Header
        header = [str(cell) if cell else "" for cell in table[0]]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * len(header)) + " |")

        # Rows
        for row in table[1:]:
            cells = [str(cell) if cell else "" for cell in row]
            # Pad se necessário
            while len(cells) < len(header):
                cells.append("")
            lines.append("| " + " | ".join(cells[:len(header)]) + " |")

        return "\n".join(lines)


# ============================================
# CAMELOT TABLE EXTRACTOR (FALLBACK 3)
# ============================================

class CamelotTableExtractor(TableExtractor):
    """
    Extrator usando Camelot.

    Modo lattice para tabelas com bordas.
    """

    name = "Camelot"

    def __init__(self):
        self._initialized = False

    def initialize(self) -> bool:
        """Verifica se Camelot está disponível."""
        if not HAS_CAMELOT:
            return False
        self._initialized = True
        return True

    def extract_tables(self, source: Any) -> list[str]:
        """Extrai tabelas de PDF usando Camelot."""
        if not self._initialized:
            return []

        tables = []
        try:
            # source é path do PDF
            extracted = camelot.read_pdf(
                str(source),
                flavor='lattice',
                pages='all'
            )
            for table in extracted:
                df = table.df
                md = df.to_markdown(index=False)
                if md:
                    tables.append(md)
        except Exception as e:
            console.print(f"[dim]Camelot error: {e}[/dim]")
        return tables


# ============================================
# TABLE CASCADE (FALLBACK AUTOMÁTICO)
# ============================================

class TableCascade:
    """
    Cascata de extração de tabelas com suporte a diferentes tipos de source.

    IMPORTANTE:
    - PyMuPDF e PDFPlumber trabalham com objetos de PÁGINA (fitz.Page)
    - Docling e Camelot trabalham com CAMINHOS DE ARQUIVO (Path/str)

    Por padrão em modo "auto", usa apenas extractors que funcionam por página
    para evitar erros. Use --table-engine docling/camelot explicitamente
    para extração de arquivo inteiro.

    Ordem (auto - per-page):
    1. PyMuPDF native (rápido)
    2. PDFPlumber (tabelas sem borda)
    """

    def __init__(
        self,
        config: ConversionConfig,
        logger: SimpleLogger,
    ):
        self.config = config
        self.logger = logger
        self.page_extractors: list[TableExtractor] = []  # Work with page objects
        self.file_extractors: list[TableExtractor] = []  # Work with file paths

    def initialize(self) -> bool:
        """Inicializa extratores na ordem de prioridade."""
        use_gpu = not self.config.force_cpu and GPU_CONFIG.is_available
        engine_preference = self.config.table_engine.lower()

        # Page-level extractors (PyMuPDF, PDFPlumber)
        page_extractors_to_try = [
            PyMuPDFTableExtractor(),
            PDFPlumberTableExtractor(),
        ]

        # File-level extractors (Docling, Camelot) - only if explicitly requested
        file_extractors_to_try = []
        if engine_preference in ("docling", "all"):
            file_extractors_to_try.append(DoclingTableExtractor(use_gpu=use_gpu))
        if engine_preference in ("camelot", "all"):
            file_extractors_to_try.append(CamelotTableExtractor())

        # Initialize page extractors
        if engine_preference in ("auto", "pymupdf", "pdfplumber", "all"):
            for extractor in page_extractors_to_try:
                if engine_preference not in ("auto", "all") and extractor.name.lower() != engine_preference:
                    continue
                try:
                    if extractor.initialize():
                        self.page_extractors.append(extractor)
                        self.logger.debug(f"Page table extractor: {extractor.name}")
                except Exception as e:
                    self.logger.debug(f"Table extractor {extractor.name} failed: {e}")

        # Initialize file extractors
        for extractor in file_extractors_to_try:
            try:
                if extractor.initialize():
                    self.file_extractors.append(extractor)
                    self.logger.debug(f"File table extractor: {extractor.name}")
            except Exception as e:
                self.logger.debug(f"Table extractor {extractor.name} failed: {e}")

        # Report status
        if self.page_extractors:
            self.logger.info(f"Tables (per-page): {self.page_extractors[0].name}")
        if self.file_extractors:
            self.logger.info(f"Tables (per-file): {self.file_extractors[0].name}")

        return bool(self.page_extractors or self.file_extractors)

    def extract_tables_from_page(self, page: Any) -> list[str]:
        """Extrai tabelas de um objeto de página (fitz.Page)."""
        for extractor in self.page_extractors:
            try:
                tables = extractor.extract_tables(page)
                if tables:
                    return tables
            except Exception as e:
                self.logger.debug(f"{extractor.name} page extraction failed: {e}")
                continue
        return []

    def extract_tables_from_file(self, file_path: Path) -> list[str]:
        """Extrai tabelas de um arquivo PDF completo."""
        for extractor in self.file_extractors:
            try:
                tables = extractor.extract_tables(file_path)
                if tables:
                    return tables
            except Exception as e:
                self.logger.debug(f"{extractor.name} file extraction failed: {e}")
                continue
        return []

    def extract_tables(self, source: Any) -> list[str]:
        """
        Extrai tabelas detectando automaticamente o tipo de source.

        - Se for Path/str: usa file_extractors ou page_extractors
        - Se for objeto de página: usa page_extractors
        """
        # Check if source is a file path
        if isinstance(source, (str, Path)):
            path = Path(source) if isinstance(source, str) else source
            if path.exists() and path.is_file():
                return self.extract_tables_from_file(path)

        # Assume it's a page object
        return self.extract_tables_from_page(source)


# ============================================
# EXCEL CONVERTER (MULTI-SHEET)
# ============================================

class ExcelConverter:
    """
    Conversor de Excel para Markdown com suporte multi-sheet.

    Features:
    - Cada aba vira uma seção ## Sheet: Nome
    - Read-only mode para eficiência
    - Suporte a .xls legado via xlrd
    """

    SUPPORTED_EXTENSIONS = {".xlsx", ".xls", ".xlsm"}

    def __init__(self, logger: SimpleLogger):
        self.logger = logger

    def can_convert(self, path: Path) -> bool:
        """Verifica se pode converter."""
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def convert(self, path: Path) -> tuple[str, int]:
        """Converte Excel para Markdown."""
        suffix = path.suffix.lower()

        if suffix == ".xls" and HAS_XLRD:
            return self._convert_xls_legacy(path)
        elif HAS_OPENPYXL:
            return self._convert_xlsx(path)
        elif HAS_PANDAS:
            return self._convert_pandas(path)
        else:
            self.logger.error("No Excel library available")
            return "", 0

    def _convert_xlsx(self, path: Path) -> tuple[str, int]:
        """Converte .xlsx usando openpyxl read-only mode."""
        try:
            wb = load_workbook(path, read_only=True, data_only=True)
            sections = []

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = list(ws.iter_rows(values_only=True))
                if not rows:
                    continue

                md_table = self._rows_to_markdown(rows, sheet_name)
                sections.append(md_table)

            wb.close()
            return "\n\n---\n\n".join(sections), len(sections)
        except Exception as e:
            self.logger.error(f"Excel conversion failed: {e}")
            return "", 0

    def _convert_xls_legacy(self, path: Path) -> tuple[str, int]:
        """Converte .xls legado usando xlrd."""
        try:
            wb = xlrd.open_workbook(str(path))
            sections = []

            for sheet_name in wb.sheet_names():
                ws = wb.sheet_by_name(sheet_name)
                rows = []
                for row_idx in range(ws.nrows):
                    row = [ws.cell_value(row_idx, col_idx) for col_idx in range(ws.ncols)]
                    rows.append(row)

                if rows:
                    md_table = self._rows_to_markdown(rows, sheet_name)
                    sections.append(md_table)

            return "\n\n---\n\n".join(sections), len(sections)
        except Exception as e:
            self.logger.error(f"XLS conversion failed: {e}")
            return "", 0

    def _convert_pandas(self, path: Path) -> tuple[str, int]:
        """Fallback usando pandas."""
        try:
            xls = pd.ExcelFile(path)
            sections = []

            for sheet_name in xls.sheet_names:
                df = xls.parse(sheet_name)
                md_table = f"## Sheet: {sheet_name}\n\n{df.to_markdown(index=False)}"
                sections.append(md_table)

            return "\n\n---\n\n".join(sections), len(sections)
        except Exception as e:
            self.logger.error(f"Pandas Excel conversion failed: {e}")
            return "", 0

    def _rows_to_markdown(self, rows: list[tuple], sheet_name: str) -> str:
        """Converte rows para GFM table."""
        lines = [f"## Sheet: {sheet_name}\n"]

        if not rows:
            return "\n".join(lines)

        # Header
        header = rows[0]
        header_cells = [str(cell) if cell is not None else "" for cell in header]
        lines.append("| " + " | ".join(header_cells) + " |")
        lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

        # Data rows
        for row in rows[1:]:
            cells = [str(cell) if cell is not None else "" for cell in row]
            while len(cells) < len(header_cells):
                cells.append("")
            lines.append("| " + " | ".join(cells[:len(header_cells)]) + " |")

        return "\n".join(lines)


# ============================================
# ZIP EXTRACTOR (5 NÍVEIS + SEGURANÇA)
# ============================================

@dataclass
class ExtractionStats:
    """Estatísticas de extração."""
    files_per_level: dict[int, int] = field(default_factory=dict)
    zips_per_level: dict[int, int] = field(default_factory=dict)
    total_files: int = 0
    total_zips: int = 0
    max_depth_reached: int = 0
    zip_bombs_detected: int = 0

    def add_file(self, level: int) -> None:
        self.files_per_level[level] = self.files_per_level.get(level, 0) + 1
        self.total_files += 1
        self.max_depth_reached = max(self.max_depth_reached, level)

    def add_zip(self, level: int) -> None:
        self.zips_per_level[level] = self.zips_per_level.get(level, 0) + 1
        self.total_zips += 1


def calculate_hash(file_path: Path) -> str:
    """Calcula SHA-256 para deduplicação."""
    try:
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            # Lê primeiro 1MB para hash rápido
            sha256.update(f.read(1024 * 1024))
        return sha256.hexdigest()
    except OSError:
        return str(file_path.absolute())


class ZIPExtractor:
    """
    Extrator de ZIP seguro com 5 níveis de profundidade.

    Features:
    - 5 níveis de aninhamento (era 3)
    - Detecção de ZIP bombs (ratio > 100x)
    - Proteção path traversal
    - SHA-256 deduplicação
    - Progress tracking por nível
    """

    def __init__(
        self,
        config: ConversionConfig,
        logger: SimpleLogger,
    ):
        self.config = config
        self.logger = logger
        self.processed_hashes: set[str] = set()
        self.stats = ExtractionStats()

    def is_zip_bomb(self, zip_file: zipfile.ZipFile) -> bool:
        """Detecta ZIP bombs por ratio de compressão."""
        try:
            total_compressed = sum(info.compress_size for info in zip_file.infolist())
            total_uncompressed = sum(info.file_size for info in zip_file.infolist())

            if total_compressed == 0:
                return False

            ratio = total_uncompressed / total_compressed
            if ratio > MAX_ZIP_RATIO:
                self.stats.zip_bombs_detected += 1
                self.logger.warning(f"ZIP bomb detected! Ratio: {ratio:.1f}x (max: {MAX_ZIP_RATIO}x)")
                return True
            return False
        except Exception:
            return False

    def is_path_traversal(self, member_path: str, base_dir: Path) -> bool:
        """Verifica tentativa de path traversal."""
        try:
            # Resolve o caminho absoluto
            full_path = (base_dir / member_path).resolve()
            # Verifica se está dentro do diretório base
            return not str(full_path).startswith(str(base_dir.resolve()))
        except Exception:
            return True

    def extract_zip(
        self,
        zip_path: Path,
        output_dir: Path,
        progress_callback: Optional[Callable] = None,
    ) -> list[Path]:
        """
        Extrai ZIP com segurança e 5 níveis de profundidade.

        Usa BFS (Breadth-First Search) para extração nível a nível.
        """
        output_dir.mkdir(parents=True, exist_ok=True)
        all_files: list[Path] = []

        # BFS queue: (zip_path, depth, extract_to, parent_name)
        queue_items: deque[tuple[Path, int, Path, str]] = deque([
            (zip_path, 0, output_dir, zip_path.stem)
        ])
        processed: set[str] = set()
        safety_counter = 0

        while queue_items and safety_counter < 10000:
            safety_counter += 1
            current_zip, depth, extract_to, parent_name = queue_items.popleft()

            # Check depth limit
            if depth > self.config.max_zip_depth:
                self.logger.warning(f"Max depth {self.config.max_zip_depth} reached, skipping")
                continue

            # Check for duplicates
            zip_hash = calculate_hash(current_zip)
            if zip_hash in processed:
                continue
            processed.add(zip_hash)

            try:
                with zipfile.ZipFile(current_zip, "r") as zf:
                    # Security: Check for ZIP bomb
                    if self.is_zip_bomb(zf):
                        self.logger.error(f"ZIP bomb detected, skipping: {current_zip.name}")
                        continue

                    for member in zf.namelist():
                        if member.endswith("/"):
                            continue

                        original_name = Path(member).name
                        if not original_name:
                            continue

                        # Security: Check path traversal
                        if self.is_path_traversal(member, extract_to):
                            self.logger.warning(f"Path traversal attempt blocked: {member}")
                            continue

                        # Generate safe filename
                        if depth > 0 and self.config.flat_names:
                            safe_name = f"L{depth}_{parent_name}_{original_name}"
                        else:
                            safe_name = original_name

                        safe_path = extract_to / safe_name

                        try:
                            with zf.open(member) as src:
                                content = src.read()

                                # Size limit
                                if len(content) > MAX_FILE_SIZE_MB * 1024 * 1024:
                                    self.logger.warning(f"File too large, skipping: {member}")
                                    continue

                                safe_path.write_bytes(content)

                            suffix = safe_path.suffix.lower()

                            # Queue nested ZIPs
                            if suffix == ".zip" and self.config.process_nested_zips:
                                if depth < self.config.max_zip_depth:
                                    nested_dir = extract_to / f"{safe_path.stem}_L{depth + 1}"
                                    nested_dir.mkdir(exist_ok=True)
                                    queue_items.append((
                                        safe_path,
                                        depth + 1,
                                        nested_dir,
                                        safe_path.stem
                                    ))
                                    self.stats.add_zip(depth)

                            # Collect processable files
                            elif suffix in SUPPORTED_EXTENSIONS or suffix in LEGACY_TO_MODERN:
                                all_files.append(safe_path)
                                self.stats.add_file(depth)

                        except Exception as e:
                            self.logger.debug(f"Error extracting {member}: {e}")

            except Exception as e:
                self.logger.error(f"Error processing {current_zip.name}: {e}")

        self.logger.increment("zips_extracted", self.stats.total_zips + 1)
        return all_files


# ============================================
# PDF CONVERTER (COMPLETO)
# ============================================

class PDFConverter:
    """
    Conversor completo de PDF para Markdown.

    Pipeline:
    1. Se use_layout_extraction → pymupdf4llm.to_markdown() (melhor qualidade)
    2. Fallback: extração página-a-página com flags aprimorados
    3. Se scanned → preprocessamento + OCR
    4. Extração de tabelas em paralelo
    5. Combinação com page markers
    """

    def __init__(
        self,
        config: ConversionConfig,
        logger: SimpleLogger,
        ocr_cascade: OCRCascade,
        table_cascade: TableCascade,
        preprocessor: ImagePreprocessor,
    ):
        self.config = config
        self.logger = logger
        self.ocr = ocr_cascade
        self.tables = table_cascade
        self.preprocessor = preprocessor

    def convert(self, pdf_path: Path) -> tuple[str, int]:
        """
        Converte PDF para Markdown usando a melhor estratégia disponível.

        Prioridade:
        1. pymupdf4llm com layout (se use_layout_extraction=True e disponível)
        2. Extração página-a-página com flags aprimorados
        3. OCR para páginas scanned
        """
        if not HAS_PYMUPDF:
            self.logger.error("PyMuPDF not available")
            return "", 0

        try:
            # Estratégia 1: pymupdf4llm.to_markdown() (melhor para layout complexo)
            if self.config.use_layout_extraction and HAS_PYMUPDF4LLM:
                result = self._convert_with_pymupdf4llm(pdf_path)
                if result[0]:  # Se teve sucesso
                    return result

            # Estratégia 2: Extração página-a-página aprimorada
            return self._convert_page_by_page(pdf_path)

        except Exception as e:
            self.logger.error(f"PDF conversion failed: {e}")
            return "", 0

    # ============================================
    # v7.2: Detecção de PPT→PDF e Otimizações
    # ============================================
    # Ref: https://github.com/pymupdf/RAG/issues/78

    def _detect_ppt_origin(self, pdf_path: Path) -> tuple[bool, str]:
        """
        Detecta se PDF foi gerado a partir de PowerPoint/Apresentação.

        Issue #78 (pymupdf/RAG): PPT→PDF frequentemente causa problemas
        com multi-column extraction devido à forma como PowerPoint
        renderiza slides para PDF.

        Args:
            pdf_path: Caminho para o arquivo PDF

        Returns:
            tuple[bool, str]: (is_ppt_origin, detected_producer)
        """
        try:
            with fitz.open(str(pdf_path)) as doc:
                metadata = doc.metadata

                # Extrair producer e creator do metadata
                producer = (metadata.get("producer", "") or "").lower().strip()
                creator = (metadata.get("creator", "") or "").lower().strip()

                # Indicadores de origem de apresentação
                ppt_indicators = [
                    "powerpoint",           # Microsoft PowerPoint
                    "microsoft office",     # Suite Office (genérico)
                    "libreoffice impress",  # LibreOffice Impress
                    "keynote",              # Apple Keynote
                    "slides",               # Google Slides
                    "prezi",                # Prezi
                    "canva",                # Canva
                    "impress",              # LibreOffice Impress (parcial)
                    "presentation",         # Genérico
                ]

                # Verificar producer
                for indicator in ppt_indicators:
                    if indicator in producer:
                        return True, f"producer: {producer}"

                # Verificar creator
                for indicator in ppt_indicators:
                    if indicator in creator:
                        return True, f"creator: {creator}"

                return False, ""

        except Exception as e:
            self.logger.debug(f"Erro ao detectar origem PPT: {e}")
            return False, ""

    def _get_ppt_optimized_params(self) -> dict[str, Any]:
        """
        Retorna parâmetros otimizados para PDFs de apresentações.

        Mitigação para Issue #78 (pymupdf/RAG):
        - PPT→PDF tende a criar layouts "falsos" com múltiplas colunas
        - Texto em caixas de texto separadas aparece como colunas
        - Gráficos e formas interferem na ordem de leitura

        Returns:
            dict: Parâmetros para pymupdf4llm.to_markdown()
        """
        return {
            "ignore_graphics": True,      # Reduz confusão de layout
            "table_strategy": "lines",    # Menos restritivo
            "force_text": True,           # Força extração de texto
        }

    def _log_ppt_detection(self, pdf_path: Path, is_ppt: bool, producer: str) -> None:
        """Registra detecção de PPT com logging apropriado."""
        if is_ppt:
            self.logger.warning(
                f"⚠️ PDF originado de apresentação detectado: {pdf_path.name}\n"
                f"   Metadata: {producer}\n"
                f"   Aplicando otimizações para Issue #78 (multi-column workaround)\n"
                f"   Ref: https://github.com/pymupdf/RAG/issues/78"
            )
            self.logger.increment("ppt_detected")
        else:
            self.logger.debug(f"PDF {pdf_path.name}: origem PPT não detectada")

    def _convert_with_pymupdf4llm(self, pdf_path: Path) -> tuple[str, int]:
        """
        Conversão layout-aware usando pymupdf4llm.

        v7.2 Features:
            - Header/footer exclusion automático (requer pymupdf-layout)
            - Detecção e otimização para PPT→PDF (Issue #78)
            - Parâmetros configuráveis via ConversionConfig
            - Logging detalhado do status e features

        Referências:
            - API: https://pymupdf.readthedocs.io/en/latest/pymupdf4llm/api.html
            - Layout: https://pymupdf.readthedocs.io/en/latest/pymupdf-layout/index.html
            - Issue #78: https://github.com/pymupdf/RAG/issues/78
        """
        try:
            # ═══════════════════════════════════════════════════════════
            # FASE 1: Logging de Status do Layout Extension
            # ═══════════════════════════════════════════════════════════
            if HAS_PYMUPDF_LAYOUT:
                self.logger.debug(
                    f"PyMuPDF Layout Extension: ATIVO\n"
                    f"  ├─ Header exclusion: {'✓ Ativo' if self.config.exclude_headers else '✗ Inativo'}\n"
                    f"  ├─ Footer exclusion: {'✓ Ativo' if self.config.exclude_footers else '✗ Inativo'}\n"
                    f"  ├─ Multi-column: ML-enhanced detection\n"
                    f"  └─ Table detection: ML-enhanced"
                )
            else:
                self.logger.debug(
                    f"PyMuPDF Layout Extension: INDISPONÍVEL\n"
                    f"  ├─ Usando heurísticas (menos preciso)\n"
                    f"  └─ Table strategy: {self.config.table_strategy}"
                )

            # ═══════════════════════════════════════════════════════════
            # FASE 2: Detecção de PPT→PDF (Issue #78)
            # ═══════════════════════════════════════════════════════════
            is_ppt_origin = False
            ppt_producer = ""

            if self.config.auto_detect_ppt_origin:
                is_ppt_origin, ppt_producer = self._detect_ppt_origin(pdf_path)
                self._log_ppt_detection(pdf_path, is_ppt_origin, ppt_producer)

            # ═══════════════════════════════════════════════════════════
            # FASE 3: Construir Parâmetros Base
            # ═══════════════════════════════════════════════════════════
            params: dict[str, Any] = {
                "doc": str(pdf_path),
                "write_images": False,
                "embed_images": False,
                "force_text": self.config.force_text_over_graphics,
                "show_progress": False,
                "page_chunks": False,
            }

            # ═══════════════════════════════════════════════════════════
            # FASE 4: Configurar Margens
            # ═══════════════════════════════════════════════════════════
            if self.config.layout_margins == (0,):
                params["margins"] = 0
            elif len(self.config.layout_margins) == 1:
                params["margins"] = self.config.layout_margins[0]
            else:
                params["margins"] = list(self.config.layout_margins)

            # ═══════════════════════════════════════════════════════════
            # FASE 5: Parâmetros que REQUEREM PyMuPDF Layout
            # ═══════════════════════════════════════════════════════════
            if HAS_PYMUPDF_LAYOUT:
                # header=False significa EXCLUIR headers
                # footer=False significa EXCLUIR footers
                params["header"] = not self.config.exclude_headers
                params["footer"] = not self.config.exclude_footers

                self.logger.debug(
                    f"Layout params: header={params['header']} (excluir={self.config.exclude_headers}), "
                    f"footer={params['footer']} (excluir={self.config.exclude_footers})"
                )
            else:
                # Sem layout extension, usar table_strategy
                params["table_strategy"] = self.config.table_strategy

            # ═══════════════════════════════════════════════════════════
            # FASE 6: Parâmetros Opcionais de Graphics/Images
            # ═══════════════════════════════════════════════════════════
            if self.config.ignore_graphics_for_text:
                params["ignore_graphics"] = True
                self.logger.debug("ignore_graphics=True: Gráficos serão ignorados")

            if self.config.ignore_images_for_text:
                params["ignore_images"] = True
                self.logger.debug("ignore_images=True: Imagens serão ignoradas")

            # ═══════════════════════════════════════════════════════════
            # FASE 7: Aplicar Otimizações PPT se Detectado
            # ═══════════════════════════════════════════════════════════
            if is_ppt_origin:
                ppt_params = self._get_ppt_optimized_params()
                params.update(ppt_params)
                self.logger.info(
                    f"Parâmetros PPT aplicados: {ppt_params}\n"
                    f"  └─ Ref: https://github.com/pymupdf/RAG/issues/78"
                )

            # ═══════════════════════════════════════════════════════════
            # FASE 8: Executar Conversão
            # ═══════════════════════════════════════════════════════════
            self.logger.debug(f"Chamando pymupdf4llm.to_markdown() com params: {params}")

            md_text = pymupdf4llm.to_markdown(**params)

            # ═══════════════════════════════════════════════════════════
            # FASE 9: Pós-processamento UTF-8
            # ═══════════════════════════════════════════════════════════
            md_text = ensure_utf8_text(md_text)

            # ═══════════════════════════════════════════════════════════
            # FASE 10: Obter Contagem de Páginas e Finalizar
            # ═══════════════════════════════════════════════════════════
            with fitz.open(str(pdf_path)) as doc:
                page_count = len(doc)

            self.logger.increment("pdfs_converted")
            self.logger.debug(
                f"pymupdf4llm converteu {page_count} páginas com sucesso\n"
                f"  ├─ Layout: {'ML-enhanced' if HAS_PYMUPDF_LAYOUT else 'Heuristic'}\n"
                f"  ├─ PPT optimized: {'Sim' if is_ppt_origin else 'Não'}\n"
                f"  └─ Output: {len(md_text)} chars"
            )

            return md_text, page_count

        except Exception as e:
            self.logger.warning(
                f"pymupdf4llm falhou: {e}\n"
                f"  └─ Usando fallback página-a-página"
            )
            return "", 0

    def _convert_page_by_page(self, pdf_path: Path) -> tuple[str, int]:
        """
        Conversão página-a-página com flags de extração aprimorados.

        Fallback robusto que funciona mesmo sem pymupdf4llm.
        """
        try:
            with fitz.open(str(pdf_path)) as doc:
                page_count = len(doc)
                pages_text = []

                # Flags de extração otimizados
                # TEXT_PRESERVE_WHITESPACE: Preserva espaçamento original
                # TEXT_PRESERVE_LIGATURES: Mantém fi, fl, etc. como caracteres únicos
                # NÃO usar TEXT_INHIBIT_SPACES (causa fragmentação!)
                extraction_flags = (
                    fitz.TEXT_PRESERVE_WHITESPACE |
                    fitz.TEXT_PRESERVE_LIGATURES
                )

                # TEXT_DEHYPHENATE: Junta palavras hifenizadas (PyMuPDF 1.24+)
                if hasattr(fitz, 'TEXT_DEHYPHENATE'):
                    extraction_flags |= fitz.TEXT_DEHYPHENATE

                for page_num, page in enumerate(doc):
                    page_content = []

                    # 1. Tenta extração de texto nativo com flags aprimorados
                    text = page.get_text("text", flags=extraction_flags, sort=True)

                    if text and len(text.strip()) > 50:
                        # PDF tem texto extraível
                        page_content.append(ensure_utf8_text(text))

                        # 2. Extrai tabelas
                        tables = self.tables.extract_tables_from_page(page)
                        for table in tables:
                            page_content.append(f"\n{table}\n")
                            self.logger.increment("tables_extracted")

                    else:
                        # PDF scanned - usar OCR
                        self.logger.debug(f"Page {page_num + 1}: Using OCR (scanned)")

                        # Renderiza página como imagem
                        pix = page.get_pixmap(dpi=self.config.preprocessing_dpi)
                        img = np.frombuffer(pix.samples, dtype=np.uint8)
                        img = img.reshape(pix.h, pix.w, pix.n)

                        # Converte para BGR se necessário
                        if pix.n == 4:  # RGBA
                            img = cv2.cvtColor(img, cv2.COLOR_RGBA2BGR)
                        elif pix.n == 1:  # Grayscale
                            img = cv2.cvtColor(img, cv2.COLOR_GRAY2BGR)

                        # 3. Preprocessamento
                        processed = self.preprocessor.preprocess(img)

                        # 4. OCR
                        ocr_text = self.ocr.process_image(processed)
                        if ocr_text:
                            page_content.append(ensure_utf8_text(ocr_text))
                            self.logger.increment("ocr_pages")

                    # Combina conteúdo da página
                    if page_content:
                        if self.config.add_page_markers:
                            header = f"\n\n---\n\n**[Página {page_num + 1} de {page_count}]**\n\n"
                            pages_text.append(header + "\n\n".join(page_content))
                        else:
                            pages_text.append("\n\n".join(page_content))

                self.logger.increment("pdfs_converted")
                return "\n".join(pages_text), page_count

        except Exception as e:
            self.logger.error(f"Page-by-page conversion failed: {e}")
            return "", 0


# ============================================
# MARKDOWN ASSEMBLER
# ============================================

class MarkdownAssembler:
    """Assembler final de Markdown com frontmatter."""

    def __init__(self, config: ConversionConfig):
        self.config = config

    def assemble(
        self,
        content: str,
        source_path: Path,
        page_count: int = 0,
        conversion_tool: str = "Batch2MD-v7.1",
    ) -> str:
        """Monta markdown final com frontmatter."""
        if not content:
            return ""

        # Normaliza
        content = ensure_utf8_text(content)

        if not self.config.add_frontmatter:
            return content

        # Frontmatter
        frontmatter = f"""---
source_format: {source_path.suffix.lower().lstrip('.')}
source_file: {source_path.name}
conversion_tool: {conversion_tool}
conversion_date: {datetime.now().isoformat()}
page_count: {page_count}
---

"""
        return frontmatter + content


# ============================================
# MAIN CONVERTER CLASS
# ============================================

class Batch2MDv7:
    """
    Conversor principal Batch2MD v7.

    Orquestra todos os componentes:
    - ZIP extraction (5 níveis)
    - OCR cascade (PaddleOCR → EasyOCR → Tesseract)
    - Table cascade (Docling → PyMuPDF → PDFPlumber → Camelot)
    - Document converters (PDF, Excel, Office, etc.)
    """

    def __init__(self, config: ConversionConfig):
        self.config = config
        self.logger = SimpleLogger(verbose=config.verbose)

        # Componentes
        self.preprocessor = ImagePreprocessor(
            target_dpi=config.preprocessing_dpi,
            enable_deskew=config.enable_deskew,
            enable_binarize=config.enable_binarize,
            enable_denoise=config.enable_denoise,
            use_gpu=not config.force_cpu,
        )

        self.ocr_cascade = OCRCascade(config, self.logger)
        self.table_cascade = TableCascade(config, self.logger)
        self.zip_extractor = ZIPExtractor(config, self.logger)
        self.excel_converter = ExcelConverter(self.logger)
        self.assembler = MarkdownAssembler(config)

        self._pdf_converter: Optional[PDFConverter] = None

    def initialize(self) -> bool:
        """Inicializa todos os componentes."""
        console.print("\n[bold cyan]═══════════════════════════════════════════════════════════════[/bold cyan]")
        console.print("[bold cyan]         Batch2MD v7.1 - GPU-Accelerated (NO LLM)              [/bold cyan]")
        console.print("[bold cyan]═══════════════════════════════════════════════════════════════[/bold cyan]")

        # Environment info
        console.print(f"\n[cyan]📍 Environment: {ENVIRONMENT}[/cyan]")

        # GPU Status (mais detalhado)
        if GPU_CONFIG.is_available:
            console.print(f"[bold green]🚀 GPU DETECTED: {GPU_CONFIG.gpu_name}[/bold green]")
            console.print(f"   [green]├─ VRAM: {GPU_CONFIG.vram_gb:.1f} GB[/green]")
            console.print(f"   [green]├─ Type: {GPU_CONFIG.gpu_type.upper()}[/green]")
            console.print(f"   [green]├─ OCR Batch Size: {GPU_CONFIG.paddle_batch_size}[/green]")
            console.print(f"   [green]└─ Memory Fraction: {GPU_CONFIG.memory_fraction:.0%}[/green]")
            if self.config.force_cpu:
                console.print("[yellow]⚠ GPU available but --cpu flag forces CPU mode[/yellow]")
        else:
            console.print(f"[yellow]⚠ NO GPU: {GPU_CONFIG.gpu_name}[/yellow]")
            console.print("[yellow]  Running in CPU mode (slower)[/yellow]")

        # ═══════════════════════════════════════════════════════════════════
        # v7.2 - Display Detalhado de Layout Features
        # ═══════════════════════════════════════════════════════════════════
        console.print(f"\n[cyan]📄 PDF Extraction:[/cyan]")
        if HAS_PYMUPDF4LLM:
            console.print(f"   [dim]├─ pymupdf4llm: ✓ Instalado[/dim]")

            if HAS_PYMUPDF_LAYOUT:
                console.print(f"   [dim]├─ pymupdf-layout: [bold green]✓ ML-ENHANCED[/bold green][/dim]")
                console.print(f"   [dim]│  ├─ Header exclusion: {'✓ Ativo' if self.config.exclude_headers else '✗ Inativo'}[/dim]")
                console.print(f"   [dim]│  ├─ Footer exclusion: {'✓ Ativo' if self.config.exclude_footers else '✗ Inativo'}[/dim]")
                console.print(f"   [dim]│  ├─ Multi-column: ✓ ML-enhanced detection[/dim]")
                console.print(f"   [dim]│  └─ Table detection: ✓ ML-enhanced[/dim]")
            else:
                console.print(f"   [dim]├─ pymupdf-layout: [yellow]✗ HEURISTIC MODE[/yellow][/dim]")
                console.print(f"   [dim]│  ├─ Header/footer exclusion: ✗ Indisponível[/dim]")
                console.print(f"   [dim]│  ├─ Multi-column: Heurísticas básicas[/dim]")
                console.print(f"   [dim]│  └─ Table strategy: {self.config.table_strategy}[/dim]")

            console.print(f"   [dim]├─ PPT detection: {'✓ Ativo' if self.config.auto_detect_ppt_origin else '✗ Inativo'}[/dim]")
            console.print(f"   [dim]└─ Layout extraction: {'✓ Enabled' if self.config.use_layout_extraction else '✗ Disabled'}[/dim]")

            if HAS_PYMUPDF_LAYOUT:
                console.print("   [green]   → ML-enhanced extraction para documentos complexos[/green]")
        else:
            console.print("   [dim]├─ pymupdf4llm: ✗ (usando fallback)[/dim]")
            console.print("   [dim]└─ Modo: Extração página-a-página[/dim]")

        # v7.2 - MPS Warning (Apple Silicon)
        if IS_MPS_SYSTEM:
            console.print(f"   [yellow]⚠️ MPS (Apple Silicon): Layout features operam em CPU[/yellow]")

        console.print("")

        # Inicializa OCR
        if not self.ocr_cascade.initialize():
            self.logger.error("Failed to initialize OCR")
            return False

        # Inicializa Tables (opcional)
        self.table_cascade.initialize()

        # PDF Converter
        self._pdf_converter = PDFConverter(
            self.config,
            self.logger,
            self.ocr_cascade,
            self.table_cascade,
            self.preprocessor,
        )

        # Preprocessing config
        console.print(f"\n[cyan]🔧 Preprocessing:[/cyan]")
        console.print(f"   [dim]├─ Deskew: {'✓' if self.config.enable_deskew else '✗'}[/dim]")
        console.print(f"   [dim]├─ Binarize: {'✓' if self.config.enable_binarize else '✗'}[/dim]")
        console.print(f"   [dim]├─ Denoise: {'✓' if self.config.enable_denoise else '✗'}[/dim]")
        console.print(f"   [dim]└─ DPI: {self.config.preprocessing_dpi}[/dim]")

        console.print("\n[bold green]✓ Initialization complete[/bold green]\n")
        return True

    def convert_file(self, file_path: Path) -> tuple[str, int]:
        """Converte um arquivo para Markdown."""
        suffix = file_path.suffix.lower()

        # PDF
        if suffix == ".pdf":
            return self._pdf_converter.convert(file_path)

        # Excel
        elif suffix in ExcelConverter.SUPPORTED_EXTENSIONS:
            return self.excel_converter.convert(file_path)

        # Office docs via MarkItDown
        elif HAS_MARKITDOWN and suffix in {".docx", ".doc", ".pptx", ".ppt"}:
            try:
                md = MarkItDown()
                result = md.convert(str(file_path))
                if result and result.text_content:
                    return ensure_utf8_text(result.text_content), 1
            except Exception as e:
                self.logger.debug(f"MarkItDown failed: {e}")

        # Email
        elif suffix == ".msg" and HAS_EXTRACT_MSG:
            try:
                msg = extract_msg.Message(str(file_path))
                content = f"# {msg.subject}\n\n"
                content += f"**From:** {msg.sender}\n"
                content += f"**To:** {msg.to}\n"
                content += f"**Date:** {msg.date}\n\n"
                content += msg.body or ""
                return ensure_utf8_text(content), 1
            except Exception as e:
                self.logger.debug(f"extract-msg failed: {e}")

        # Images via OCR
        elif suffix in {".jpg", ".jpeg", ".png", ".tiff", ".bmp", ".gif", ".webp"}:
            try:
                if HAS_OPENCV:
                    img = cv2.imread(str(file_path))
                    if img is not None:
                        processed = self.preprocessor.preprocess(img)
                        text = self.ocr_cascade.process_image(processed)
                        if text:
                            return ensure_utf8_text(text), 1
            except Exception as e:
                self.logger.debug(f"Image OCR failed: {e}")

        # HTML files - convert to markdown
        elif suffix in {".html", ".htm"}:
            try:
                # Try MarkItDown first (best HTML to MD conversion)
                if HAS_MARKITDOWN:
                    md = MarkItDown()
                    result = md.convert(str(file_path))
                    if result and result.text_content:
                        return ensure_utf8_text(result.text_content), 1

                # Fallback: basic HTML to text extraction
                html_content = file_path.read_text(encoding="utf-8", errors="replace")
                # Remove script and style tags
                html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
                html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
                # Convert common HTML elements to markdown
                html_content = re.sub(r'<br\s*/?>', '\n', html_content, flags=re.IGNORECASE)
                html_content = re.sub(r'<p[^>]*>', '\n\n', html_content, flags=re.IGNORECASE)
                html_content = re.sub(r'</p>', '', html_content, flags=re.IGNORECASE)
                html_content = re.sub(r'<h1[^>]*>(.*?)</h1>', r'\n# \1\n', html_content, flags=re.DOTALL | re.IGNORECASE)
                html_content = re.sub(r'<h2[^>]*>(.*?)</h2>', r'\n## \1\n', html_content, flags=re.DOTALL | re.IGNORECASE)
                html_content = re.sub(r'<h3[^>]*>(.*?)</h3>', r'\n### \1\n', html_content, flags=re.DOTALL | re.IGNORECASE)
                html_content = re.sub(r'<h4[^>]*>(.*?)</h4>', r'\n#### \1\n', html_content, flags=re.DOTALL | re.IGNORECASE)
                html_content = re.sub(r'<li[^>]*>(.*?)</li>', r'- \1\n', html_content, flags=re.DOTALL | re.IGNORECASE)
                html_content = re.sub(r'<[au]l[^>]*>', '\n', html_content, flags=re.IGNORECASE)
                html_content = re.sub(r'</[au]l>', '\n', html_content, flags=re.IGNORECASE)
                html_content = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', html_content, flags=re.DOTALL | re.IGNORECASE)
                html_content = re.sub(r'<b[^>]*>(.*?)</b>', r'**\1**', html_content, flags=re.DOTALL | re.IGNORECASE)
                html_content = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', html_content, flags=re.DOTALL | re.IGNORECASE)
                html_content = re.sub(r'<i[^>]*>(.*?)</i>', r'*\1*', html_content, flags=re.DOTALL | re.IGNORECASE)
                html_content = re.sub(r'<a[^>]*href="([^"]*)"[^>]*>(.*?)</a>', r'[\2](\1)', html_content, flags=re.DOTALL | re.IGNORECASE)
                # Remove remaining HTML tags
                html_content = re.sub(r'<[^>]+>', '', html_content)
                # Decode HTML entities
                html_content = html_module.unescape(html_content)
                # Clean up whitespace
                html_content = re.sub(r'\n{3,}', '\n\n', html_content)
                html_content = html_content.strip()
                return ensure_utf8_text(html_content), 1
            except Exception as e:
                self.logger.debug(f"HTML conversion failed: {e}")

        # Plain text (excluding HTML)
        elif suffix in {".txt", ".md", ".xml", ".json", ".csv"}:
            try:
                text = file_path.read_text(encoding="utf-8", errors="replace")
                return ensure_utf8_text(text), 1
            except Exception as e:
                self.logger.debug(f"Text read failed: {e}")

        # Fallback to MarkItDown
        if HAS_MARKITDOWN:
            try:
                md = MarkItDown()
                result = md.convert(str(file_path))
                if result and result.text_content:
                    return ensure_utf8_text(result.text_content), 1
            except Exception:
                pass

        self.logger.warning(f"Could not convert: {file_path.name}")
        return "", 0

    def process(
        self,
        input_path: Path,
        output_dir: Path,
    ) -> dict:
        """
        Processa input (ZIP, pasta ou arquivo) e gera Markdown.

        Returns:
            Dict com estatísticas de processamento.
        """
        output_dir.mkdir(parents=True, exist_ok=True)
        files_to_process: list[Path] = []
        temp_dir: Optional[Path] = None

        try:
            # Determina tipo de input
            if input_path.suffix.lower() == ".zip":
                # Extrai ZIP
                temp_dir = Path(tempfile.mkdtemp(prefix="batch2md_"))
                self.logger.info(f"Extracting ZIP (up to {self.config.max_zip_depth} levels)...")
                files_to_process = self.zip_extractor.extract_zip(input_path, temp_dir)
                self.logger.success(f"Extracted {len(files_to_process)} files")

            elif input_path.is_dir():
                # Lista arquivos recursivamente
                if self.config.recursive:
                    files_to_process = [
                        f for f in input_path.rglob("*")
                        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
                    ]
                else:
                    files_to_process = [
                        f for f in input_path.iterdir()
                        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
                    ]

            elif input_path.is_file():
                files_to_process = [input_path]

            else:
                self.logger.error(f"Invalid input: {input_path}")
                return {"error": "Invalid input"}

            # Dry run
            if self.config.dry_run:
                self.logger.info(f"Dry run: {len(files_to_process)} files found")
                for f in files_to_process[:10]:
                    console.print(f"  - {f.name}")
                if len(files_to_process) > 10:
                    console.print(f"  ... and {len(files_to_process) - 10} more")
                return {"files_found": len(files_to_process)}

            # Processa arquivos
            total = len(files_to_process)
            converted = 0
            errors = 0

            with Progress(
                SpinnerColumn(),
                TextColumn("[progress.description]{task.description}"),
                BarColumn(),
                TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
                TimeElapsedColumn(),
                console=console,
            ) as progress:
                task = progress.add_task("Converting...", total=total)

                for file_path in files_to_process:
                    try:
                        # Converte
                        content, page_count = self.convert_file(file_path)

                        if content:
                            # Assembla com frontmatter
                            final_md = self.assembler.assemble(
                                content,
                                file_path,
                                page_count,
                            )

                            # Salva
                            output_name = file_path.stem + ".md"
                            output_path = output_dir / output_name
                            output_path.write_text(final_md, encoding="utf-8")
                            converted += 1
                            self.logger.increment("files_processed")
                        else:
                            errors += 1

                    except Exception as e:
                        self.logger.error(f"Error processing {file_path.name}: {e}")
                        errors += 1

                    progress.advance(task)

            # Estatísticas
            stats = {
                "total_files": total,
                "converted": converted,
                "errors": errors,
                "ocr_pages": self.logger.counters.get("ocr_pages", 0),
                "tables_extracted": self.logger.counters.get("tables_extracted", 0),
                "zip_stats": {
                    "total_zips": self.zip_extractor.stats.total_zips,
                    "max_depth": self.zip_extractor.stats.max_depth_reached,
                    "zip_bombs_blocked": self.zip_extractor.stats.zip_bombs_detected,
                },
            }

            # Sumário
            console.print(f"\n[bold green]═══ Conversion Complete ═══[/bold green]")
            console.print(f"[green]✓ Converted: {converted}/{total} files[/green]")
            if stats["ocr_pages"] > 0:
                console.print(f"[cyan]  OCR pages: {stats['ocr_pages']}[/cyan]")
            if stats["tables_extracted"] > 0:
                console.print(f"[cyan]  Tables extracted: {stats['tables_extracted']}[/cyan]")
            if errors > 0:
                console.print(f"[yellow]  Errors: {errors}[/yellow]")

            return stats

        finally:
            # Cleanup temp
            if temp_dir and temp_dir.exists():
                try:
                    shutil.rmtree(temp_dir)
                except Exception:
                    pass


# ============================================
# CLI
# ============================================

@app.command()
def main(
    input_path: Path = typer.Argument(
        ...,
        help="Input ZIP, file, or folder",
        exists=True,
    ),
    output_dir: Optional[Path] = typer.Argument(
        None,
        help="Output directory (or use --out)",
    ),
    out_option: Path = typer.Option(
        Path("./_md_output"),
        "--out", "-o",
        help="Output directory for markdown files",
    ),
    # Processing
    workers: int = typer.Option(
        GPU_CONFIG.preprocessing_workers,  # Dinâmico baseado na GPU
        "--workers", "-w",
        min=1, max=64,
        help=f"Number of parallel workers (auto: {GPU_CONFIG.preprocessing_workers} for {GPU_CONFIG.gpu_type})",
    ),
    recursive: bool = typer.Option(
        True,
        "--recursive/--no-recursive",
        help="Process directories recursively",
    ),
    # ZIP
    nested_zips: bool = typer.Option(
        True,
        "--nested-zips/--no-nested-zips",
        help="Process nested ZIP files",
    ),
    max_depth: int = typer.Option(
        5,
        "--max-depth",
        min=1, max=10,
        help="Maximum ZIP nesting depth",
    ),
    # OCR
    ocr_engine: str = typer.Option(
        "auto",
        "--ocr-engine",
        help="OCR engine: auto, paddle, easyocr, tesseract",
    ),
    ocr_lang: str = typer.Option(
        "pt,en",
        "--ocr-lang",
        help="OCR languages (comma-separated)",
    ),
    # Preprocessing
    deskew: bool = typer.Option(
        True,
        "--deskew/--no-deskew",
        help="Enable deskewing",
    ),
    binarize: bool = typer.Option(
        True,
        "--binarize/--no-binarize",
        help="Enable Sauvola binarization",
    ),
    denoise: bool = typer.Option(
        True,
        "--denoise/--no-denoise",
        help="Enable noise reduction",
    ),
    dpi: int = typer.Option(
        300,
        "--dpi",
        min=72, max=600,
        help="DPI for image preprocessing",
    ),
    # Tables
    table_engine: str = typer.Option(
        "auto",
        "--table-engine",
        help="Table extractor: auto, docling, pymupdf, pdfplumber, camelot",
    ),
    # Output
    frontmatter: bool = typer.Option(
        True,
        "--frontmatter/--no-frontmatter",
        help="Add YAML frontmatter",
    ),
    page_markers: bool = typer.Option(
        True,
        "--page-markers/--no-page-markers",
        help="Add page markers to output",
    ),
    # GPU
    cpu: bool = typer.Option(
        False,
        "--cpu",
        help="Force CPU mode (disable GPU)",
    ),
    force_gpu: bool = typer.Option(
        False,
        "--force-gpu",
        help="Require GPU (fail if not available)",
    ),
    # ═══════════════════════════════════════════════════════════════════
    # v7.2 - Layout Options
    # ═══════════════════════════════════════════════════════════════════
    exclude_headers: bool = typer.Option(
        True,
        "--exclude-headers/--include-headers",
        help="Excluir headers (requer pymupdf-layout). Default: excluir.",
    ),
    exclude_footers: bool = typer.Option(
        True,
        "--exclude-footers/--include-footers",
        help="Excluir footers (requer pymupdf-layout). Default: excluir.",
    ),
    ignore_graphics: bool = typer.Option(
        False,
        "--ignore-graphics",
        help="Ignorar gráficos para melhor extração de texto.",
    ),
    ignore_images: bool = typer.Option(
        False,
        "--ignore-images",
        help="Ignorar imagens para melhor extração de texto.",
    ),
    table_strategy: str = typer.Option(
        "lines_strict",
        "--table-strategy",
        help="Estratégia de tabelas: lines_strict, lines, text.",
    ),
    no_ppt_detection: bool = typer.Option(
        False,
        "--no-ppt-detection",
        help="Desabilitar detecção automática de PPT→PDF.",
    ),
    # Debug
    verbose: bool = typer.Option(
        False,
        "--verbose", "-v",
        help="Verbose output",
    ),
    dry_run: bool = typer.Option(
        False,
        "--dry-run",
        help="List files without converting",
    ),
):
    """
    Batch2MD v7.2 - GPU-Accelerated Document Converter (NO LLM)

    Converts documents to Markdown using traditional OCR (PaddleOCR, EasyOCR, Tesseract)
    instead of LLMs. Optimized for Google Colab A100/V100/T4.

    v7.2 adds header/footer exclusion, PPT→PDF detection with optimizations (Issue #78),
    and MPS (Apple Silicon) detection with CPU fallback.

    v7.1 adds pymupdf-layout integration for better layout-aware text extraction,
    fixing character fragmentation issues (especially with numbers and accented chars).

    Examples:

        # Basic usage
        python Batch2MD-v7.py input.zip --out ./markdown_output

        # Force PaddleOCR with Portuguese
        python Batch2MD-v7.py docs/ --ocr-engine paddle --ocr-lang pt,en

        # v7.2: Include headers and footers (by default they are excluded)
        python Batch2MD-v7.py document.pdf --include-headers --include-footers

        # v7.2: Ignore graphics for cleaner text extraction
        python Batch2MD-v7.py slides.pdf --ignore-graphics

        # CPU-only mode
        python Batch2MD-v7.py document.pdf --cpu --ocr-engine tesseract

        # Maximum quality preprocessing
        python Batch2MD-v7.py scanned.pdf --deskew --binarize --denoise --dpi 300
    """
    # Resolve output directory (positional argument takes precedence over --out)
    final_output_dir = output_dir if output_dir is not None else out_option

    # Validate GPU requirements
    if force_gpu and not GPU_CONFIG.is_available:
        console.print("[bold red]ERROR: --force-gpu specified but no GPU available![/bold red]")
        console.print(f"[red]GPU Status: {GPU_CONFIG.gpu_name}[/red]")
        console.print("\n[yellow]Para usar GPU, certifique-se de:[/yellow]")
        console.print("[yellow]  1. CUDA Toolkit instalado[/yellow]")
        console.print("[yellow]  2. paddlepaddle-gpu (não paddlepaddle)[/yellow]")
        console.print("[yellow]  3. PyTorch com CUDA support[/yellow]")
        raise typer.Exit(code=1)

    if force_gpu and cpu:
        console.print("[bold red]ERROR: Cannot use --force-gpu and --cpu together![/bold red]")
        raise typer.Exit(code=1)

    # Build config
    config = ConversionConfig(
        recursive=recursive,
        workers=workers,
        process_nested_zips=nested_zips,
        max_zip_depth=max_depth,
        ocr_engine=ocr_engine,
        ocr_languages=ocr_lang.split(","),
        enable_deskew=deskew,
        enable_binarize=binarize,
        enable_denoise=denoise,
        preprocessing_dpi=dpi,
        table_engine=table_engine,
        add_frontmatter=frontmatter,
        add_page_markers=page_markers,
        force_cpu=cpu,
        force_gpu=force_gpu,
        # v7.2 - Layout options
        exclude_headers=exclude_headers,
        exclude_footers=exclude_footers,
        ignore_graphics_for_text=ignore_graphics,
        ignore_images_for_text=ignore_images,
        table_strategy=table_strategy,
        auto_detect_ppt_origin=not no_ppt_detection,
        verbose=verbose,
        dry_run=dry_run,
    )

    # Create and run converter
    converter = Batch2MDv7(config)

    if not converter.initialize():
        raise typer.Exit(code=1)

    try:
        stats = converter.process(input_path, final_output_dir)

        if stats.get("error"):
            raise typer.Exit(code=1)

    except KeyboardInterrupt:
        console.print("\n[yellow]Cancelled by user[/yellow]")
        raise typer.Exit(code=130)
    except Exception as e:
        console.print(f"\n[red]Error: {e}[/red]")
        if verbose:
            import traceback
            traceback.print_exc()
        raise typer.Exit(code=1)


if __name__ == "__main__":
    app()
